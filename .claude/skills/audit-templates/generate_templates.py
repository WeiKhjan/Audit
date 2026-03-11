#!/usr/bin/env python3
"""
Audit Template .docx Generator
Generates well-formatted Word documents for audit letters and confirmations (T1-T16).
Reads master_data.json for variable substitution.

Usage:
    python generate_templates.py all
    python generate_templates.py T1 T4 T8
    python generate_templates.py T8 --directors "Cho Chin Lai"
    python generate_templates.py T4 --banks "Public Bank Berhad"
"""

import json
import os
import re
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# ============================================================
# CONFIGURATION
# ============================================================

TEMPLATE_MAP = {
    "T1":  {"wp": "A1", "folder": "A_Planning",             "filename": "A1_Engagement_Letter.docx",              "title": "Engagement Letter",               "isa": "ISA 210", "per_instance": False},
    "T2":  {"wp": "H1", "folder": "H_Engagement_Tracker",   "filename": "H1_PBC_Request_Letter.docx",             "title": "PBC Request Letter",              "isa": "",        "per_instance": False},
    "T3":  {"wp": "",   "folder": "F_Completion",            "filename": "F_Management_Letter.docx",               "title": "Management Letter",               "isa": "ISA 265", "per_instance": False},
    "T4":  {"wp": "C9", "folder": "C_Assets",                "filename": "C9_Bank_Confirmation_{bank}.docx",       "title": "Bank Confirmation",               "isa": "ISA 505", "per_instance": "bank"},
    "T5":  {"wp": "C9", "folder": "C_Assets",                "filename": "C9_Bank_Authorization_{bank}.docx",      "title": "Bank Authorization Letter",       "isa": "ISA 505", "per_instance": "bank"},
    "T6":  {"wp": "C6", "folder": "C_Assets",                "filename": "C6_Debtor_Confirmation_{name}.docx",     "title": "Debtor/Borrower Confirmation",    "isa": "ISA 505", "per_instance": "debtor"},
    "T7":  {"wp": "D8", "folder": "D_Liabilities_Equity",    "filename": "D8_Creditor_Confirmation_{name}.docx",   "title": "Creditor Confirmation",           "isa": "ISA 505", "per_instance": "creditor"},
    "T8":  {"wp": "D9", "folder": "D_Liabilities_Equity",    "filename": "D9_Director_Confirmation_{director}.docx","title": "Director Confirmation",           "isa": "ISA 550", "per_instance": "director"},
    "T9":  {"wp": "",   "folder": "F_Completion",            "filename": "F_Legal_Confirmation.docx",              "title": "Legal Confirmation",              "isa": "ISA 501", "per_instance": False},
    "T10": {"wp": "C5", "folder": "C_Assets",                "filename": "C5_Stock_Confirmation.docx",             "title": "Stock Confirmation",              "isa": "ISA 505", "per_instance": False},
    "T11": {"wp": "F7", "folder": "F_Completion",            "filename": "F7_Management_Representation.docx",      "title": "Management Representation Letter","isa": "ISA 580", "per_instance": False},
    "T12": {"wp": "F1", "folder": "F_Completion",            "filename": "F1_Director_Support_Letter.docx",        "title": "Director Support Letter",         "isa": "ISA 570", "per_instance": False},
    "T13": {"wp": "F6", "folder": "F_Completion",            "filename": "F6_Audit_Adjustments.docx",              "title": "Summary of Audit Adjustments",    "isa": "ISA 450", "per_instance": False},
    "T14": {"wp": "F6", "folder": "F_Completion",            "filename": "F6_Uncorrected_Differences.docx",        "title": "Uncorrected Differences",         "isa": "ISA 450", "per_instance": False},
    "T15": {"wp": "F3", "folder": "F_Completion",            "filename": "F3_Director_Remuneration_{director}.docx","title": "Directors' Remuneration Confirmation","isa": "CA 2016 s.230","per_instance": "director"},
    "T16": {"wp": "F3", "folder": "F_Completion",            "filename": "F3_Director_Shareholding_{director}.docx","title": "Directors' Shareholding Confirmation","isa": "CA 2016 s.59","per_instance": "director"},
}

TODAY = datetime.now().strftime("%d %B %Y")
TODAY_SHORT = datetime.now().strftime("%Y-%m-%d")


# ============================================================
# HELPERS
# ============================================================

def load_master_data():
    """Load and flatten master_data.json into key:value dict."""
    path = os.path.join(os.getcwd(), "master_data.json")
    if not os.path.exists(path):
        print(f"ERROR: master_data.json not found in {os.getcwd()}")
        sys.exit(1)
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    flat = {}
    for key, obj in data.get("variables", {}).items():
        val = obj.get("value", "")
        if isinstance(val, (int, float)):
            fmt = obj.get("format", "")
            if fmt in ("currency", "currency_bracket"):
                if val < 0:
                    flat[key] = f"(RM {abs(val):,.2f})"
                else:
                    flat[key] = f"RM {val:,.2f}"
            else:
                flat[key] = str(val)
        else:
            flat[key] = str(val)
    return flat


def v(vars_dict, key, default="[___]"):
    """Get variable value with fallback."""
    return vars_dict.get(key, default)


def safe_filename(name):
    """Convert name to safe filename component."""
    return re.sub(r'[^A-Za-z0-9]', '_', name).strip('_')


def abbrev_name(name):
    """Abbreviate name: 'Cho Chin Lai' -> 'CCL'."""
    parts = name.split()
    return "".join(p[0].upper() for p in parts if p)


def fmt_rm(value):
    """Format RM value for display."""
    if isinstance(value, str):
        return value
    if value < 0:
        return f"(RM {abs(value):,.2f})"
    return f"RM {value:,.2f}"


# ============================================================
# DOCUMENT FORMATTING HELPERS
# ============================================================

def set_doc_margins(doc):
    """Set standard letter margins."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_letterhead(doc, vars_dict):
    """Add firm letterhead to document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(v(vars_dict, "audit_firm"))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(v(vars_dict, "audit_firm_address"))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Tel: {v(vars_dict, 'audit_firm_tel')}  |  Email: {v(vars_dict, 'audit_firm_email')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Horizontal line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1E40AF"/></w:pBdr>')
    pPr.append(pBdr)


def add_company_letterhead(doc, vars_dict):
    """Add company letterhead (for letters FROM the company)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(v(vars_dict, "company_name"))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"({v(vars_dict, 'company_reg_no')})")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(v(vars_dict, "registered_address"))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1E40AF"/></w:pBdr>')
    pPr.append(pBdr)


def add_date(doc, date_str=None):
    """Add date line."""
    p = doc.add_paragraph()
    p.add_run(date_str or TODAY).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)


def add_addressee(doc, lines):
    """Add addressee block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10)
        if line.startswith("**") or line == lines[0]:
            run.bold = True


def add_subject(doc, text):
    """Add RE: subject line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"RE: {text}")
    run.bold = True
    run.underline = True
    run.font.size = Pt(10)


def add_heading(doc, text, level=1):
    """Add heading with styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return h


def add_para(doc, text, bold=False, italic=False, size=10, space_after=6):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_numbered_item(doc, number, text, indent=0.5):
    """Add a numbered paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"({number}) ")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E40AF"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        for c, text in enumerate(row_data):
            cell = row.cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_signature_block(doc, name, designation="Director", nric="", include_stamp=False):
    """Add a signature block."""
    doc.add_paragraph()  # spacing
    p = add_para(doc, "Signature: ____________________________", size=10)
    p.paragraph_format.space_after = Pt(4)
    add_para(doc, f"Name: {name}", bold=True, size=10, space_after=2)
    if nric:
        add_para(doc, f"NRIC No.: {nric}", size=10, space_after=2)
    add_para(doc, f"Designation: {designation}", size=10, space_after=2)
    add_para(doc, "Date: ____________________________", size=10, space_after=12)
    if include_stamp:
        add_para(doc, "Company Stamp: ____________________________", size=10, space_after=12)


def add_footer_isa(doc, isa_ref, description):
    """Add ISA reference footer."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="1" w:color="94A3B8"/></w:pBdr>')
    pPr.append(pBdr)
    run = p.add_run(f"Prepared in accordance with {isa_ref} — {description}")
    run.font.size = Pt(8)
    run.italic = True
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)


def add_horizontal_line(doc):
    """Add a horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="CBD5E1"/></w:pBdr>')
    pPr.append(pBdr)


def add_return_instruction(doc, vars_dict):
    """Add standard return instruction block."""
    add_horizontal_line(doc)
    add_para(doc, "Please return this confirmation directly to our auditors at:", bold=True, size=10)
    add_para(doc, v(vars_dict, "audit_firm"), bold=True, size=10, space_after=2)
    add_para(doc, v(vars_dict, "audit_firm_address"), size=10, space_after=2)
    add_para(doc, f"(Do NOT return this letter to {v(vars_dict, 'company_name')}.)", italic=True, size=9, space_after=12)


# ============================================================
# TEMPLATE BUILDERS (T1-T16)
# ============================================================

def build_T1(doc, vd):
    """T1 — Engagement Letter (ISA 210)"""
    add_letterhead(doc, vd)
    add_date(doc)
    add_addressee(doc, [
        "The Board of Directors",
        v(vd, "company_name"),
        f"({v(vd, 'company_reg_no')})",
        v(vd, "registered_address"),
    ])
    add_para(doc, "Dear Directors,")
    add_subject(doc, f"AUDIT OF THE FINANCIAL STATEMENTS OF {v(vd, 'company_name')} FOR THE FINANCIAL YEAR ENDED {v(vd, 'year_end_date')}")

    add_para(doc, f"We are pleased to confirm our acceptance and understanding of the terms and conditions of our engagement to audit the financial statements of {v(vd, 'company_name')} (Company No. {v(vd, 'company_reg_no')}) for the financial year ended {v(vd, 'year_end_date')}.")

    add_heading(doc, "1. OBJECTIVE AND SCOPE OF THE AUDIT", level=2)
    add_para(doc, "The objective of our audit is to obtain reasonable assurance about whether the financial statements as a whole are free from material misstatement, whether due to fraud or error, and to issue an auditor's report that includes our opinion.")
    add_para(doc, "The audit will be conducted in accordance with:")
    add_para(doc, "- Section 266 of the Companies Act 2016 — statutory requirement for annual audit;", size=10, space_after=2)
    add_para(doc, "- International Standards on Auditing (ISA) as adopted in Malaysia by the MIA;", size=10, space_after=2)
    add_para(doc, f"- {v(vd, 'reporting_framework')} as the applicable financial reporting framework.", size=10)

    add_heading(doc, "2. AUDITOR'S RESPONSIBILITIES", level=2)
    responsibilities = [
        "Plan and perform the audit to obtain reasonable assurance about whether the financial statements are free from material misstatement;",
        "Identify and assess the risks of material misstatement and design audit procedures responsive to those risks;",
        "Obtain an understanding of internal control relevant to the audit;",
        "Evaluate the appropriateness of accounting policies and reasonableness of accounting estimates;",
        "Conclude on the appropriateness of the going concern basis of accounting;",
        "Evaluate the overall presentation, structure, and content of the financial statements;",
        "Communicate significant audit findings and deficiencies in internal control (ISA 265).",
    ]
    for i, resp in enumerate(responsibilities):
        add_numbered_item(doc, chr(97 + i), resp)
    add_para(doc, "Reasonable assurance is a high level of assurance but is not a guarantee that an audit will always detect a material misstatement when it exists.")

    add_heading(doc, "3. DIRECTORS' RESPONSIBILITIES", level=2)
    dir_resp = [
        f"The preparation of financial statements that give a true and fair view in accordance with {v(vd, 'reporting_framework')};",
        "Such internal control as the directors determine is necessary to enable preparation of financial statements free from material misstatement;",
        "Providing us with access to all relevant information, records, and documentation;",
        "Providing us with additional information that we may request for the purpose of the audit;",
        "Providing us with unrestricted access to persons within the entity from whom we determine it necessary to obtain audit evidence;",
        "Informing us of all known instances of non-compliance with laws and regulations;",
        "Providing us with a written Management Representation Letter prior to the issuance of the audit report.",
    ]
    for i, resp in enumerate(dir_resp):
        add_numbered_item(doc, chr(97 + i), resp)

    add_heading(doc, "4. REPORTING FRAMEWORK", level=2)
    add_para(doc, f"The financial statements will be prepared in accordance with the {v(vd, 'reporting_framework')} issued by the Malaysian Accounting Standards Board (MASB).")

    add_heading(doc, "5. INHERENT LIMITATIONS OF AN AUDIT", level=2)
    add_para(doc, "Because of the inherent limitations of an audit, together with the inherent limitations of internal control, there is an unavoidable risk that some material misstatements may not be detected, even though the audit is properly planned and performed in accordance with ISA.")

    add_heading(doc, "6. ACCESS TO RECORDS", level=2)
    add_para(doc, "We shall be given full and complete access at all times to all books, records, documents, accounts, vouchers, correspondence, contracts, and any other information and explanations we consider necessary for the purpose of the audit, in accordance with Section 269 of the Companies Act 2016.")

    add_heading(doc, "7. FEE ARRANGEMENT", level=2)
    add_para(doc, f"Our professional fee for the statutory audit for the financial year ended {v(vd, 'year_end_date')} is estimated at {v(vd, 'audit_fee')}, exclusive of out-of-pocket expenses and service tax where applicable.")

    add_heading(doc, "8. TERMS AND CONDITIONS", level=2)
    add_para(doc, "This engagement letter shall be effective for the audit of the financial statements and shall continue in force for subsequent audit engagements unless terminated or superseded by a revised engagement letter.")

    add_heading(doc, "9. CONFIDENTIALITY", level=2)
    add_para(doc, "All information obtained during the course of our audit is confidential and shall not be disclosed to any third party without prior written consent, except where required by law, professional standards, or regulatory authorities.")

    add_heading(doc, "10. ACCEPTANCE AND AGREEMENT", level=2)
    add_para(doc, "Please sign and return the enclosed copy of this letter to indicate your acknowledgement and agreement to the terms and conditions of the audit engagement.")

    add_horizontal_line(doc)
    add_para(doc, f"FOR AND ON BEHALF OF {v(vd, 'audit_firm')}:", bold=True)
    add_signature_block(doc, v(vd, "engagement_partner"), "Engagement Partner")

    add_horizontal_line(doc)
    add_para(doc, f"ACKNOWLEDGED AND AGREED ON BEHALF OF {v(vd, 'company_name')}:", bold=True)
    add_para(doc, "Director 1:", bold=True)
    add_signature_block(doc, v(vd, "director_1_name"))
    add_para(doc, "Director 2:", bold=True)
    add_signature_block(doc, v(vd, "director_2_name"))

    add_footer_isa(doc, "ISA 210", "Agreeing the Terms of Audit Engagements")


def build_T2(doc, vd):
    """T2 — PBC Request Letter"""
    add_letterhead(doc, vd)
    add_date(doc)
    add_addressee(doc, [
        "The Board of Directors",
        v(vd, "company_name"),
        f"({v(vd, 'company_reg_no')})",
        v(vd, "registered_address"),
    ])
    add_para(doc, "Dear Directors,")
    add_subject(doc, f"AUDIT OF FINANCIAL STATEMENTS FOR THE FINANCIAL YEAR ENDED {v(vd, 'year_end_date')} — REQUEST FOR DOCUMENTS AND INFORMATION (PBC LIST)")
    add_para(doc, f"We refer to our engagement to audit the financial statements of {v(vd, 'company_name')} for the financial year ended {v(vd, 'year_end_date')}.")
    add_para(doc, f"To enable us to carry out the audit efficiently, we kindly request that the following documents and information be prepared and made available to us by {v(vd, 'pbc_deadline_date')}.")

    add_heading(doc, "A. CRITICAL ITEMS", level=2)
    critical = [
        ["1", "Signed Engagement Letter (T1)", "T1"],
        ["2", "Director balance confirmations", "C2/D2"],
        ["3", "Director support letter — written undertaking for financial support", "F1"],
        ["4", "Bank confirmation authorization letter (T5)", "T5"],
        ["5", "Signed Management Representation Letter (MRL)", "F7"],
        ["6", "Complete borrower listing and loan agreements", "C1"],
        ["7", "Directors' account general ledger", "C2/D2"],
        ["8", "Financial statements approval — directors' resolution", "F8"],
    ]
    add_table(doc, ["No.", "Document / Information", "Ref"], critical, [0.5, 5.0, 0.7])

    add_heading(doc, "B. STATUTORY & CORPORATE", level=2)
    statutory = [
        ["9", "SSM company profile / e-Info printout (latest)", "A2"],
        ["10", f"Board resolutions and circular resolutions for FY {v(vd, 'fye_year')}", "A2"],
        ["11", f"AGM / EGM minutes for FY {v(vd, 'fye_year')}", "A2"],
        ["12", "Valid money lending licence", "A2"],
        ["13", "Directors' declaration of related party transactions", "F3"],
    ]
    add_table(doc, ["No.", "Document / Information", "Ref"], statutory, [0.5, 5.0, 0.7])

    add_heading(doc, "C. FINANCIAL RECORDS", level=2)
    financial = [
        ["14", f"Trial balance as at {v(vd, 'year_end_date')}", "A3"],
        ["15", f"General ledger for FY {v(vd, 'fye_year')}", "A3"],
        ["16", f"Balance sheet as at {v(vd, 'year_end_date')}", "A3"],
        ["17", f"Profit and loss for FY {v(vd, 'fye_year')}", "A3"],
        ["18", f"Fixed asset register as at {v(vd, 'year_end_date')}", "B1"],
        ["19", "Prior year audited financial statements", "A3"],
        ["20", "Prior year tax computation", "E1"],
        ["21", "Prior year audit working papers", "A3"],
    ]
    add_table(doc, ["No.", "Document / Information", "Ref"], financial, [0.5, 5.0, 0.7])

    add_heading(doc, "D. BANK & CASH", level=2)
    bank = [
        ["22", f"Bank statements for ALL accounts for December {v(vd, 'fye_year')}", "B2"],
        ["23", "Bank statements for subsequent period (January)", "B2"],
        ["24", f"Bank reconciliation statements as at {v(vd, 'year_end_date')}", "B2"],
    ]
    add_table(doc, ["No.", "Document / Information", "Ref"], bank, [0.5, 5.0, 0.7])

    add_heading(doc, "DELIVERY INSTRUCTIONS", level=2)
    add_para(doc, f"Deadline: Please ensure all documents are submitted by {v(vd, 'pbc_deadline_date')}.", bold=True)
    add_para(doc, "Documents may be provided in physical copy or electronic format (PDF preferred for scanned documents, Excel for schedules).")

    add_horizontal_line(doc)
    add_para(doc, "Yours faithfully,")
    add_para(doc, v(vd, "audit_firm"), bold=True)
    add_signature_block(doc, v(vd, "engagement_partner"), "Engagement Partner")


def build_T3(doc, vd):
    """T3 — Management Letter (ISA 265)"""
    add_letterhead(doc, vd)
    add_para(doc, "PRIVATE AND CONFIDENTIAL", bold=True, size=10)
    add_date(doc)
    add_addressee(doc, [
        "The Board of Directors",
        v(vd, "company_name"),
        f"({v(vd, 'company_reg_no')})",
        v(vd, "registered_address"),
    ])
    add_para(doc, "Dear Directors,")
    add_subject(doc, f"COMMUNICATION OF SIGNIFICANT DEFICIENCIES IN INTERNAL CONTROL — FYE {v(vd, 'year_end_date')}")

    add_heading(doc, "1. INTRODUCTION", level=2)
    add_para(doc, f"In accordance with ISA 265, we wish to bring to your attention certain matters that came to our notice during the course of our statutory audit of the financial statements of {v(vd, 'company_name')} for the financial year ended {v(vd, 'year_end_date')}.")

    add_heading(doc, "2. SUMMARY OF FINDINGS", level=2)
    add_table(doc, ["No.", "Observation", "Risk Rating", "Recommendation", "Management Response", "Target Date"], [
        ["1", "[To be completed]", "[H/M/L]", "[To be completed]", "[To be completed]", "[Date]"],
        ["2", "[To be completed]", "[H/M/L]", "[To be completed]", "[To be completed]", "[Date]"],
    ])

    add_heading(doc, "3. CONCLUSION", level=2)
    add_para(doc, "The responsibility for maintaining an adequate system of internal control and for the prevention and detection of irregularities and fraud rests with the directors of the Company.")
    add_para(doc, "We would be pleased to discuss any of the above matters with you at your convenience.")

    add_horizontal_line(doc)
    add_para(doc, "Yours faithfully,")
    add_para(doc, v(vd, "audit_firm"), bold=True)
    add_signature_block(doc, v(vd, "engagement_partner"), "Engagement Partner")

    add_horizontal_line(doc)
    add_para(doc, "ACKNOWLEDGEMENT BY THE BOARD OF DIRECTORS:", bold=True)
    add_signature_block(doc, v(vd, "director_1_name"))
    add_signature_block(doc, v(vd, "director_2_name"))

    add_footer_isa(doc, "ISA 265", "Communicating Deficiencies in Internal Control")


def build_T4(doc, vd):
    """T4 — Bank Confirmation Request (ISA 505)"""
    add_letterhead(doc, vd)
    add_date(doc)
    add_addressee(doc, [
        "The Branch Manager",
        v(vd, "bank_name"),
        v(vd, "bank_branch"),
        v(vd, "bank_branch_address"),
    ])
    add_para(doc, "Dear Sir / Madam,")
    add_subject(doc, f"REQUEST FOR BANK CONFIRMATION — {v(vd, 'company_name')} ({v(vd, 'company_reg_no')}) AS AT {v(vd, 'year_end_date')}")
    add_para(doc, f"In connection with our statutory audit of the financial statements of {v(vd, 'company_name')} for the financial year ended {v(vd, 'year_end_date')}, we would be grateful if you could confirm directly to us the following information.")

    add_heading(doc, "A. BANK ACCOUNTS", level=2)
    add_para(doc, "Please confirm the following details for ALL accounts held in the name of the Company:")
    add_table(doc, ["No.", "Account Number", "Currency", "Balance"], [
        ["1", v(vd, "bank_account_1"), "MYR", v(vd, "cash_bank_pbb_3194091531")],
        ["2", v(vd, "bank_account_2"), "MYR", v(vd, "cash_bank_pbb_3203516834")],
    ])

    for section, title in [("B", "FIXED DEPOSITS"), ("C", "OUTSTANDING FACILITIES"),
                           ("D", "GUARANTEES AND LETTERS OF CREDIT"), ("E", "SECURITIES HELD OR PLEDGED"),
                           ("F", "AUTHORIZED SIGNATORIES")]:
        add_heading(doc, f"{section}. {title}", level=2)
        add_para(doc, f"Please confirm details of all {title.lower()} as at {v(vd, 'year_end_date')}.")

    add_heading(doc, "BANK'S CONFIRMATION", level=2)
    add_para(doc, "We confirm that the information provided above is correct / the correct information is as follows:")
    add_para(doc, "____________________________________________________________")
    add_para(doc, "____________________________________________________________")
    add_signature_block(doc, "____________________________", "Bank Officer")

    add_horizontal_line(doc)
    add_para(doc, "Yours faithfully,")
    add_para(doc, v(vd, "audit_firm"), bold=True)
    add_signature_block(doc, v(vd, "engagement_partner"), "Engagement Partner")

    add_para(doc, "IMPORTANT: Please send your reply directly to the auditors and NOT to the Company. This is a requirement under ISA 505.", bold=True, italic=True, size=9)
    add_footer_isa(doc, "ISA 505", "External Confirmations")


def build_T5(doc, vd):
    """T5 — Bank Authorization Letter (ISA 505)"""
    add_company_letterhead(doc, vd)
    add_date(doc)
    add_addressee(doc, [
        "The Branch Manager",
        v(vd, "bank_name"),
        v(vd, "bank_branch"),
        v(vd, "bank_branch_address"),
    ])
    add_para(doc, "Dear Sir / Madam,")
    add_subject(doc, f"AUTHORIZATION TO RELEASE BANK INFORMATION TO AUDITORS — {v(vd, 'company_name')} ({v(vd, 'company_reg_no')})")
    add_para(doc, f"We, the undersigned directors of {v(vd, 'company_name')}, hereby authorize and request {v(vd, 'bank_name')} to provide all information directly to our appointed auditors.")

    add_heading(doc, "AUDITOR DETAILS", level=2)
    add_table(doc, ["", ""], [
        ["Audit Firm:", v(vd, "audit_firm")],
        ["Address:", v(vd, "audit_firm_address")],
        ["Engagement Partner:", v(vd, "engagement_partner")],
        ["Contact:", v(vd, "audit_firm_tel")],
        ["Email:", v(vd, "audit_firm_email")],
    ])

    add_heading(doc, "BANK ACCOUNTS", level=2)
    add_table(doc, ["No.", "Account Number", "Account Type"], [
        ["1", v(vd, "bank_account_1"), "Current / Savings"],
        ["2", v(vd, "bank_account_2"), "Current / Savings"],
    ])

    add_heading(doc, "SCOPE OF AUTHORIZATION", level=2)
    scope_items = [
        "Account balances as at any date requested;",
        "Bank statements for any period within the financial year;",
        "Fixed deposit details including principal, interest rates, maturity dates;",
        "Credit facility details including terms and outstanding balances;",
        "Guarantee and contingent liability details;",
        "Security and collateral details;",
        "Authorized signatory details and signing mandates;",
        "Any other information reasonably requested for the statutory audit.",
    ]
    for i, item in enumerate(scope_items):
        add_numbered_item(doc, i + 1, item)

    add_heading(doc, "VALIDITY PERIOD", level=2)
    add_para(doc, f"This authorization is valid from {v(vd, 'fy_start')} to {v(vd, 'year_end_date')}, and shall remain effective until the completion of the audit.")

    add_horizontal_line(doc)
    add_para(doc, f"FOR AND ON BEHALF OF {v(vd, 'company_name')}", bold=True)
    add_para(doc, "Director 1:", bold=True)
    add_signature_block(doc, v(vd, "director_1_name"), "Director", v(vd, "director_1_ic"), include_stamp=False)
    add_para(doc, "Director 2:", bold=True)
    add_signature_block(doc, v(vd, "director_2_name"), "Director", v(vd, "director_2_ic"), include_stamp=True)

    add_footer_isa(doc, "ISA 505", "External Confirmations")


def build_T6(doc, vd, debtor_name="[Borrower Name]"):
    """T6 — Debtor/Borrower Confirmation (ISA 505)"""
    add_letterhead(doc, vd)
    add_date(doc)
    add_addressee(doc, [debtor_name, "[Borrower Address]"])
    add_para(doc, "Dear Sir/Madam,")
    add_subject(doc, "REQUEST FOR CONFIRMATION OF LOAN BALANCE")
    add_para(doc, f"In connection with the statutory audit of {v(vd, 'company_name')} ({v(vd, 'company_reg_no')}) for the financial year ended {v(vd, 'year_end_date')}, we would be grateful if you could confirm the balance(s) due from you as at {v(vd, 'year_end_date')}.")

    add_table(doc, ["Item", "Loan Ref", "Principal (RM)", "Accrued Interest (RM)", "Total Due (RM)"], [
        ["1", "[Loan Ref]", "[Amount]", "[Amount]", "[Amount]"],
    ])

    add_heading(doc, "POSITIVE CONFIRMATION", level=2)
    add_para(doc, "This is a positive confirmation request. Please confirm whether you agree or disagree with the balance(s) above.")

    add_heading(doc, "REPLY SECTION", level=2)
    add_para(doc, "Option A: I confirm that the balance(s) stated above are correct.")
    add_para(doc, "Option B: I do NOT agree. Details of differences:")
    add_para(doc, "____________________________________________________________")
    add_signature_block(doc, "____________________________")

    add_return_instruction(doc, vd)
    add_footer_isa(doc, "ISA 505", "External Confirmations")


def build_T7(doc, vd, creditor_name="[Creditor Name]"):
    """T7 — Creditor Confirmation (ISA 505)"""
    add_letterhead(doc, vd)
    add_date(doc)
    add_addressee(doc, [creditor_name, "[Creditor Address]"])
    add_para(doc, "Dear Sir/Madam,")
    add_subject(doc, "REQUEST FOR CONFIRMATION OF ACCOUNT BALANCE")
    add_para(doc, f"In connection with the statutory audit of {v(vd, 'company_name')} for the financial year ended {v(vd, 'year_end_date')}, we would be grateful if you could confirm the balance owed to you.")

    add_table(doc, ["Description", "Amount (RM)"], [
        ["Trade payables", "[Amount]"],
        ["Accruals", "[Amount]"],
        ["Total balance", "[Total]"],
    ])

    add_heading(doc, "NEGATIVE CONFIRMATION", level=2)
    add_para(doc, "If the above balance agrees with your records, no reply is necessary.")
    add_para(doc, "If the balance does NOT agree, please complete the section below and return this letter directly to our auditors.")

    add_return_instruction(doc, vd)
    add_footer_isa(doc, "ISA 505", "External Confirmations")


def build_T8(doc, vd, director_num=1):
    """T8 — Director/Related Party Confirmation (ISA 550)"""
    dname = v(vd, f"director_{director_num}_name")
    dic = v(vd, f"director_{director_num}_ic")

    add_letterhead(doc, vd)
    add_para(doc, f"{v(vd, 'company_name')} ({v(vd, 'company_reg_no')})", bold=True)
    add_para(doc, f"Financial Year Ended: {v(vd, 'year_end_date')}", bold=True)
    add_date(doc)
    add_addressee(doc, [dname])
    add_para(doc, f"Dear {dname},")
    add_para(doc, f"In connection with the statutory audit of {v(vd, 'company_name')} for the financial year ended {v(vd, 'year_end_date')}, we request that you confirm the following information.")

    add_heading(doc, "Part A — Balance Confirmation", level=2)
    add_para(doc, f"Please confirm the balance(s) between yourself and the Company as at {v(vd, 'year_end_date')}:")
    add_table(doc, ["Description", "Amount (RM)"], [
        ["Amount due to directors", v(vd, "amount_due_directors")],
        ["Amount due from directors", "[___]"],
        ["Net balance", "[___]"],
    ])
    add_para(doc, "Do you agree with the balance(s) above?")
    add_para(doc, "[ ] Yes, I agree.    [ ] No, I disagree. Details: ________________________")

    add_heading(doc, "Part B — Transaction Confirmation", level=2)
    add_para(doc, f"Please confirm all transactions between yourself (or your related parties) and the Company during {v(vd, 'fy_start')} to {v(vd, 'year_end_date')}:")
    add_table(doc, ["Transaction Type", "Date(s)", "Amount (RM)", "Description"], [
        ["Advances to Company", "[Dates]", "[Amount]", "[Description]"],
        ["Repayments received", "[Dates]", "[Amount]", "[Description]"],
        ["Directors' fees", "[Dates]", "[Amount]", "[Description]"],
        ["Directors' remuneration", "[Dates]", "[Amount]", "[Description]"],
        ["Benefits-in-kind", "[Dates]", "[Amount]", "[Description]"],
    ])

    add_heading(doc, "Part C — Related Party Declarations", level=2)
    add_para(doc, "Please declare all related party relationships as defined under MPERS Section 33:", bold=True)
    add_table(doc, ["Related Party Name", "Relationship", "Nature"], [
        ["[Name]", "[Relationship]", "[Nature]"],
    ])
    add_para(doc, "Please declare any interests in contracts with the Company:", bold=True)
    add_para(doc, "[ ] I have no interests in any contracts with the Company.")

    add_heading(doc, "Part D — Completeness Assertion", level=2)
    add_para(doc, f"I confirm that other than as disclosed above, there are no other transactions, balances, or arrangements between myself (or my related parties) and the Company during the financial year ended {v(vd, 'year_end_date')}.")
    add_signature_block(doc, dname, "Director", dic)

    add_return_instruction(doc, vd)
    add_footer_isa(doc, "ISA 550", "Related Parties")


def build_T9(doc, vd):
    """T9 — Legal Confirmation (ISA 501)"""
    add_letterhead(doc, vd)
    add_date(doc)
    add_addressee(doc, [
        v(vd, "lawyer_firm"),
        v(vd, "lawyer_address"),
    ])
    add_para(doc, "Dear Sirs,")
    add_subject(doc, "REQUEST FOR INFORMATION — STATUTORY AUDIT")
    add_para(doc, f"In connection with the statutory audit of {v(vd, 'company_name')} ({v(vd, 'company_reg_no')}) for the financial year ended {v(vd, 'year_end_date')}, we request the following information.")

    add_heading(doc, "1. Pending or Threatened Litigation", level=2)
    add_para(doc, f"Please provide details of all pending or threatened litigation as at {v(vd, 'year_end_date')}:")
    add_table(doc, ["Case/Matter", "Nature of Claim", "Amount (RM)", "Status", "Your Assessment"], [
        ["[Case ref]", "[Nature]", "[Amount]", "[Status]", "[Assessment]"],
    ])
    add_para(doc, "[ ] There are no pending or threatened litigation matters.")

    add_heading(doc, "2. Claims and Assessments", level=2)
    add_para(doc, "Please provide details of any claims or assessments against the Company:")
    add_para(doc, "[ ] There are no claims or assessments.")

    add_heading(doc, "3. Unasserted Claims", level=2)
    add_para(doc, "[ ] No unasserted claims identified.")

    add_heading(doc, "4. Outstanding Legal Fees", level=2)
    add_table(doc, ["Description", "Period", "Amount Outstanding (RM)"], [
        ["[Services]", "[Period]", "[Amount]"],
    ])

    add_heading(doc, "5. Guarantees and Indemnities", level=2)
    add_para(doc, "[ ] No guarantees or indemnities identified.")

    add_return_instruction(doc, vd)
    add_signature_block(doc, v(vd, "engagement_partner"), "Engagement Partner")
    add_footer_isa(doc, "ISA 501", "Audit Evidence — Specific Considerations")


def build_T10(doc, vd):
    """T10 — Stock Confirmation (ISA 505)"""
    add_letterhead(doc, vd)
    add_para(doc, f"NOTE: NOT APPLICABLE for {v(vd, 'company_name')} (money lending business — no physical inventory). This template is included for completeness.", bold=True, italic=True, size=9)
    add_date(doc)
    add_addressee(doc, ["[Warehouse/Custodian Name]", "[Custodian Address]"])
    add_para(doc, "Dear Sir/Madam,")
    add_subject(doc, "REQUEST FOR CONFIRMATION OF INVENTORY HELD")
    add_para(doc, f"In connection with the statutory audit of {v(vd, 'company_name')} for the financial year ended {v(vd, 'year_end_date')}, we would be grateful if you could confirm details of any inventory held on behalf of the Company.")

    add_heading(doc, "1. Quantity and Description of Goods Held", level=2)
    add_table(doc, ["Item No.", "Description", "Unit", "Quantity", "Location"], [
        ["[No.]", "[Description]", "[Unit]", "[Qty]", "[Location]"],
    ])
    add_para(doc, "[ ] No inventory is held on behalf of the Company.")

    add_heading(doc, "2. Condition of Goods", level=2)
    add_para(doc, "Please indicate if any goods are damaged, obsolete, or subject to deterioration.")

    add_heading(doc, "3. Liens, Charges, or Encumbrances", level=2)
    add_para(doc, "[ ] No liens, charges, or encumbrances exist.")

    add_return_instruction(doc, vd)
    add_footer_isa(doc, "ISA 505", "External Confirmations")


def build_T11(doc, vd):
    """T11 — Management Representation Letter (ISA 580)"""
    doc.add_paragraph()  # spacing
    add_para(doc, "MANAGEMENT REPRESENTATION LETTER", bold=True, size=14)
    add_horizontal_line(doc)

    add_para(doc, f"From: The Directors of {v(vd, 'company_name')} ({v(vd, 'company_reg_no')})", bold=True)
    add_para(doc, f"To: {v(vd, 'audit_firm')}", bold=True)
    add_para(doc, "Date: [Date — same date as or after the audit report date]", bold=True)
    add_para(doc, f"Re: Audit of financial statements for the year ended {v(vd, 'year_end_date')}", bold=True)
    add_horizontal_line(doc)

    add_para(doc, "Dear Sirs,")
    add_para(doc, f"This representation letter is provided in connection with your audit of the financial statements of {v(vd, 'company_name')} for the year ended {v(vd, 'year_end_date')}, for the purpose of expressing an opinion as to whether the financial statements give a true and fair view in accordance with the {v(vd, 'reporting_framework')} and the Companies Act 2016.")
    add_para(doc, "We confirm that, to the best of our knowledge and belief:")

    sections = [
        ("(A) Financial Statements (ISA 580.10-11)", [
            f"We have fulfilled our responsibilities for the preparation of the financial statements in accordance with the {v(vd, 'reporting_framework')}.",
            f"The financial statements give a true and fair view of the financial position as at {v(vd, 'year_end_date')}.",
            "Significant assumptions used in making accounting estimates are reasonable and appropriate.",
        ]),
        ("(B) Completeness of Information (ISA 580.11)", [
            "All transactions have been recorded and reflected in the financial statements.",
            "We have provided you with access to all relevant information, records, and documentation.",
            "We have provided you with all minutes of meetings of shareholders and directors.",
            "There are no side agreements or undisclosed transactions.",
        ]),
        ("(C) Fraud (ISA 240.39)", [
            "We acknowledge our responsibility for internal controls to prevent and detect fraud.",
            "We have disclosed all information relating to fraud or suspected fraud.",
            "We have disclosed all allegations of fraud communicated by employees or others.",
        ]),
        ("(D) Laws & Regulations (ISA 250)", [
            "We have disclosed all known instances of non-compliance with laws and regulations.",
            "The Company has complied with all provisions of the Companies Act 2016.",
            "EPF, SOCSO, and EIS contributions are up to date.",
            "The Company's money lending licence is valid and all conditions complied with.",
        ]),
        ("(E) Accounting Estimates (ISA 540)", [
            "Assumptions and methods used in accounting estimates are reasonable.",
            "Depreciation rates and useful lives are appropriate.",
            f"Impairment assessment of receivables reflects incurred losses under {v(vd, 'reporting_framework')}.",
        ]),
        ("(F) Related Parties (ISA 550)", [
            "We have disclosed all related party relationships and transactions.",
            "All related party transactions have been on arm's length terms, or properly disclosed.",
            f"Related parties identified: {v(vd, 'director_1_name')} — Director and shareholder; {v(vd, 'director_2_name')} — Director and shareholder.",
        ]),
        ("(G) Subsequent Events (ISA 560)", [
            f"All events subsequent to {v(vd, 'year_end_date')} requiring adjustment or disclosure have been properly addressed.",
        ]),
        ("(H) Going Concern (ISA 570)", [
            f"We acknowledge the capital deficiency of {v(vd, 'total_equity')} as at {v(vd, 'year_end_date')}.",
            f"Accumulated losses of {v(vd, 'accumulated_losses_cf')} exceed share capital of {v(vd, 'share_capital')}.",
            "We believe preparation on a going concern basis is appropriate given continued director financial support.",
            f"We undertake not to demand repayment of {v(vd, 'amount_due_directors')} for at least 12 months from approval of financial statements.",
        ]),
    ]

    for heading, items in sections:
        add_heading(doc, heading, level=2)
        for i, item in enumerate(items):
            num = sum(len(s[1]) for s in sections[:sections.index((heading, items))]) + i + 1
            add_numbered_item(doc, num, item)

    add_heading(doc, "(J) Other Matters", level=2)
    add_para(doc, "There are no known litigation, claims, or assessments not disclosed in the financial statements.")
    add_para(doc, f"All bank accounts have been disclosed: {v(vd, 'bank_account_1')}, {v(vd, 'bank_account_2')}.")

    add_horizontal_line(doc)
    add_para(doc, "Yours faithfully,")
    add_para(doc, f"For and on behalf of {v(vd, 'company_name')}", bold=True)
    add_signature_block(doc, v(vd, "director_1_name"), "Director", v(vd, "director_1_ic"))
    add_signature_block(doc, v(vd, "director_2_name"), "Director", v(vd, "director_2_ic"))


def build_T12(doc, vd):
    """T12 — Director Support Letter (ISA 570)"""
    add_para(doc, "DIRECTOR'S UNDERTAKING TO PROVIDE FINANCIAL SUPPORT", bold=True, size=14)
    add_horizontal_line(doc)

    add_para(doc, f"From: The Directors of {v(vd, 'company_name')} ({v(vd, 'company_reg_no')})", bold=True)
    add_para(doc, f"To: {v(vd, 'audit_firm')}", bold=True)
    add_para(doc, "Date: [Date]", bold=True)
    add_para(doc, f"Re: Director's Undertaking to Provide Financial Support — FYE {v(vd, 'year_end_date')}", bold=True)
    add_horizontal_line(doc)

    add_para(doc, "Dear Sirs,")
    add_para(doc, f"We, the undersigned directors of {v(vd, 'company_name')} ({v(vd, 'company_reg_no')}), refer to the audit of the financial statements for the financial year ended {v(vd, 'year_end_date')}.")

    add_heading(doc, "1. Acknowledgment of Going Concern Issues", level=2)
    gc_items = [
        f"The Company is in a capital deficiency position of {v(vd, 'total_equity')} as at {v(vd, 'year_end_date')};",
        f"Accumulated losses of {v(vd, 'accumulated_losses_cf')} exceed share capital of {v(vd, 'share_capital')};",
        f"Net current liability position — current liabilities of {v(vd, 'total_current_liabilities')} exceed current assets of {v(vd, 'total_current_assets')};",
        f"Net loss of {v(vd, 'profit_after_tax')} for the financial year;",
        "Revenue has declined from the prior year.",
    ]
    for i, item in enumerate(gc_items):
        add_numbered_item(doc, chr(97 + i), item)

    add_heading(doc, "2. Undertaking to Provide Financial Support", level=2)
    undertakings = [
        "We will continue to provide financial support to the Company to enable it to continue as a going concern;",
        f"We will not demand repayment of the amount due to directors of {v(vd, 'amount_due_directors')} for at least 12 months from the date of approval of the financial statements;",
        "We will provide such additional financial support as may be required.",
    ]
    for i, item in enumerate(undertakings):
        add_numbered_item(doc, chr(97 + i), item)

    add_heading(doc, "3. Confirmation of Ability", level=2)
    add_para(doc, "We confirm that we have adequate personal financial resources to fulfill the undertaking given above for at least 12 months from the date of approval of the financial statements.")

    add_horizontal_line(doc)
    add_para(doc, "Yours faithfully,")
    add_signature_block(doc, v(vd, "director_1_name"), "Director", v(vd, "director_1_ic"))
    add_signature_block(doc, v(vd, "director_2_name"), "Director", v(vd, "director_2_ic"))


def build_T13(doc, vd):
    """T13 — Summary of Audit Adjustments (ISA 450)"""
    add_para(doc, "SUMMARY OF AUDIT ADJUSTMENTS", bold=True, size=14)
    add_horizontal_line(doc)

    add_para(doc, f"Client: {v(vd, 'company_name')} ({v(vd, 'company_reg_no')})", bold=True)
    add_para(doc, f"Year End: {v(vd, 'year_end_date')}")
    add_para(doc, f"Prepared by: {v(vd, 'engagement_partner')}  |  Date: {TODAY}")

    add_heading(doc, "1. Introduction", level=2)
    add_para(doc, f"The following audit adjustments have been identified during our audit and agreed with the directors for posting to the accounting records.")

    add_heading(doc, "2. Schedule of Agreed Audit Adjustments", level=2)
    add_table(doc, ["Ref", "Description", "DR Account", "CR Account", "DR (RM)", "CR (RM)", "Impact on Profit"], [
        ["AJ1", "Interest in arrears reclassification", "Other Payables (D8)", "Other Receivables (C7)", v(vd, "interest_in_arrears"), v(vd, "interest_in_arrears"), "Nil (reclass)"],
        ["AJ2", "[To be completed]", "", "", "", "", ""],
    ])

    add_heading(doc, "3. Total Impact Summary", level=2)
    add_table(doc, ["", "Amount (RM)"], [
        ["Total adjustments — Debit", v(vd, "interest_in_arrears")],
        ["Total adjustments — Credit", v(vd, "interest_in_arrears")],
        ["Net impact on profit before tax", "Nil"],
        ["Net impact on net assets", "Nil"],
    ])

    add_heading(doc, "5. Director Acknowledgment", level=2)
    add_para(doc, f"We acknowledge and agree to the above audit adjustments for FYE {v(vd, 'year_end_date')}.")
    add_signature_block(doc, v(vd, "director_1_name"), "Director")
    add_signature_block(doc, v(vd, "director_2_name"), "Director")


def build_T14(doc, vd):
    """T14 — Summary of Uncorrected Differences (ISA 450)"""
    add_para(doc, "SUMMARY OF UNCORRECTED AUDIT DIFFERENCES", bold=True, size=14)
    add_horizontal_line(doc)

    add_para(doc, f"Client: {v(vd, 'company_name')} ({v(vd, 'company_reg_no')})", bold=True)
    add_para(doc, f"Year End: {v(vd, 'year_end_date')}")
    add_para(doc, f"Prepared by: {v(vd, 'engagement_partner')}  |  Date: {TODAY}")
    add_para(doc, f"Planning Materiality: {v(vd, 'planning_materiality')}  |  Performance Materiality: {v(vd, 'performance_materiality')}  |  Trivial Threshold: {v(vd, 'trivial_threshold')}")

    add_heading(doc, "1. Schedule of Uncorrected Misstatements", level=2)
    add_table(doc, ["Ref", "Description", "DR (RM)", "CR (RM)", "Impact on Profit"], [
        ["PAJ1", "Professional fee accrual potentially overstated", "Up to 4,541", "Up to 4,541", "Understated up to 4,541"],
        ["PAJ2", "EPF undercontribution", "~1,000", "~1,000", "Overstated ~1,000"],
        ["PAJ3", "PY tax provision write-back", "4,141", "4,141", "Understated 4,141"],
    ])

    add_heading(doc, "2. Aggregate Effect", level=2)
    add_table(doc, ["", "Amount (RM)"], [
        ["Total potential overstatement of profit", "~1,000 (PAJ2)"],
        ["Total potential understatement of profit", "~8,682 (PAJ1+PAJ3)"],
        ["Net potential impact on profit", "Understated by ~7,682"],
    ])

    add_heading(doc, "3. Comparison to Materiality", level=2)
    add_para(doc, f"Aggregate uncorrected misstatements (~RM 7,682) EXCEED planning materiality ({v(vd, 'planning_materiality')}).")
    add_para(doc, "Note: These are POTENTIAL differences pending PBC verification.", bold=True, italic=True)

    add_heading(doc, "4. ISA 450.14 — Written Representation", level=2)
    add_para(doc, "We believe the effects of these uncorrected misstatements are immaterial to the financial statements as a whole.")

    add_heading(doc, "5. Director Acknowledgment", level=2)
    add_signature_block(doc, v(vd, "director_1_name"), "Director")
    add_signature_block(doc, v(vd, "director_2_name"), "Director")


def build_T15(doc, vd, director_num=1):
    """T15 — Directors' Remuneration Confirmation"""
    dname = v(vd, f"director_{director_num}_name")
    dic = v(vd, f"director_{director_num}_ic")

    add_para(doc, "DIRECTORS' REMUNERATION CONFIRMATION", bold=True, size=14)
    add_horizontal_line(doc)
    add_para(doc, f"{v(vd, 'company_name')} ({v(vd, 'company_reg_no')})", bold=True)
    add_para(doc, f"Financial Year Ended: {v(vd, 'year_end_date')}", bold=True)

    add_para(doc, f"Confirmation by: {dname}", bold=True, size=12)
    add_para(doc, f"To: {v(vd, 'audit_firm')}")
    add_para(doc, f"In connection with the statutory audit for FYE {v(vd, 'year_end_date')}, I confirm the following remuneration received:")

    add_heading(doc, "A. Directors' Fees & Remuneration", level=2)
    prefix = f"director_{director_num}"
    add_table(doc, ["Component", "Amount (RM)"], [
        ["Directors' fees", "[___]"],
        ["Salary", "[___]"],
        ["Bonus", "[___]"],
        ["Allowances", "[___]"],
        ["Commission", "[___]"],
        ["Total fees & remuneration", "[___]"],
    ])

    add_heading(doc, "B. Defined Contribution Plans (Employer's Portion)", level=2)
    add_table(doc, ["Component", "Amount (RM)"], [
        ["EPF — employer's contribution", "[___]"],
        ["SOCSO — employer's contribution", "[___]"],
        ["EIS — employer's contribution", "[___]"],
        ["Total defined contribution", "[___]"],
    ])

    add_heading(doc, "C. Benefits-in-Kind", level=2)
    add_table(doc, ["Component", "Amount (RM)"], [
        ["Motor vehicle", "[___]"],
        ["Fuel / petrol", "[___]"],
        ["Housing / accommodation", "[___]"],
        ["Others", "[___]"],
        ["Total benefits-in-kind", "[___]"],
    ])

    add_heading(doc, "D. Confirmation", level=2)
    add_para(doc, f"I confirm the above is a complete and accurate record of all remuneration received for FYE {v(vd, 'year_end_date')}.")
    add_para(doc, "No other remuneration, benefits, or compensation was received other than as disclosed above.")
    add_signature_block(doc, dname, "Director", dic)
    add_return_instruction(doc, vd)


def build_T16(doc, vd, director_num=1):
    """T16 — Directors' Shareholding Confirmation"""
    dname = v(vd, f"director_{director_num}_name")
    dic = v(vd, f"director_{director_num}_ic")
    shares = v(vd, f"director_{director_num}_shares")

    add_para(doc, "DIRECTORS' SHAREHOLDING CONFIRMATION", bold=True, size=14)
    add_horizontal_line(doc)
    add_para(doc, f"{v(vd, 'company_name')} ({v(vd, 'company_reg_no')})", bold=True)
    add_para(doc, f"Financial Year Ended: {v(vd, 'year_end_date')}", bold=True)

    add_para(doc, f"Confirmation by: {dname}", bold=True, size=12)
    add_para(doc, f"To: {v(vd, 'audit_firm')}")

    add_heading(doc, f"A. Direct Shareholding in {v(vd, 'company_name')}", level=2)
    add_table(doc, ["", "No. of Ordinary Shares"], [
        [f"Balance as at beginning of year ({v(vd, 'fy_start')})", shares],
        ["Shares acquired during the year", "Nil"],
        ["Shares disposed during the year", "Nil"],
        [f"Balance as at end of year ({v(vd, 'year_end_date')})", shares],
    ])

    add_heading(doc, "B. Indirect / Deemed Shareholding", level=2)
    add_para(doc, "[ ] I have no indirect or deemed interest in shares of the Company.")

    add_heading(doc, "C. Shares in Related Corporations", level=2)
    add_para(doc, "[ ] I have no interest in shares of any related corporation.")

    add_heading(doc, "D. Changes During the Financial Year", level=2)
    add_para(doc, "[ ] There were no changes in my shareholding during the financial year.")

    add_heading(doc, "E. Confirmation", level=2)
    add_para(doc, f"I confirm the above is a complete and accurate record of all my direct and indirect interests in shares as at {v(vd, 'year_end_date')}.")
    add_signature_block(doc, dname, "Director", dic)
    add_return_instruction(doc, vd)


# ============================================================
# MAIN GENERATION LOGIC
# ============================================================

def get_directors(vd):
    """Get list of directors from master data."""
    directors = []
    for i in range(1, 10):
        name = vd.get(f"director_{i}_name")
        if name and name != "[___]":
            directors.append({"num": i, "name": name, "abbrev": abbrev_name(name)})
        else:
            break
    return directors


def generate_template(ref, vd, base_dir="."):
    """Generate one or more .docx files for a template reference."""
    config = TEMPLATE_MAP.get(ref)
    if not config:
        print(f"ERROR: Unknown template reference: {ref}")
        return []

    results = []
    folder_path = os.path.join(base_dir, config["folder"])
    os.makedirs(folder_path, exist_ok=True)

    # Determine instances
    instances = [{}]  # default: single instance
    per = config["per_instance"]

    if per == "director":
        directors = get_directors(vd)
        instances = [{"director_num": d["num"], "director_name": d["name"], "abbrev": d["abbrev"]} for d in directors]
    elif per == "bank":
        # For now, single bank (Public Bank Berhad)
        bank_name = vd.get("bank_name", "Public Bank Berhad")
        instances = [{"bank_name": bank_name, "abbrev": safe_filename(bank_name)[:10]}]
    elif per == "debtor":
        instances = [{"debtor_name": "[Borrower Name]", "abbrev": "Template"}]
    elif per == "creditor":
        instances = [{"creditor_name": "[Creditor Name]", "abbrev": "Template"}]

    for inst in instances:
        doc = Document()
        set_doc_margins(doc)

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        # Build filename
        filename = config["filename"]
        if per == "director":
            filename = filename.replace("{director}", inst["abbrev"])
        elif per == "bank":
            filename = filename.replace("{bank}", inst["abbrev"])
        elif per == "debtor":
            filename = filename.replace("{name}", inst["abbrev"])
        elif per == "creditor":
            filename = filename.replace("{name}", inst["abbrev"])

        # Build document content
        builder_map = {
            "T1": lambda: build_T1(doc, vd),
            "T2": lambda: build_T2(doc, vd),
            "T3": lambda: build_T3(doc, vd),
            "T4": lambda: build_T4(doc, vd),
            "T5": lambda: build_T5(doc, vd),
            "T6": lambda: build_T6(doc, vd, inst.get("debtor_name", "[Borrower Name]")),
            "T7": lambda: build_T7(doc, vd, inst.get("creditor_name", "[Creditor Name]")),
            "T8": lambda: build_T8(doc, vd, inst.get("director_num", 1)),
            "T9": lambda: build_T9(doc, vd),
            "T10": lambda: build_T10(doc, vd),
            "T11": lambda: build_T11(doc, vd),
            "T12": lambda: build_T12(doc, vd),
            "T13": lambda: build_T13(doc, vd),
            "T14": lambda: build_T14(doc, vd),
            "T15": lambda: build_T15(doc, vd, inst.get("director_num", 1)),
            "T16": lambda: build_T16(doc, vd, inst.get("director_num", 1)),
        }

        builder = builder_map.get(ref)
        if builder:
            builder()
        else:
            print(f"WARNING: No builder for {ref}")
            continue

        # Save
        filepath = os.path.join(folder_path, filename)
        doc.save(filepath)

        # Build ATTACHMENT comment
        title = config["title"]
        if per == "director":
            title += f" - {inst['director_name']}"
        elif per == "bank":
            title += f" - {inst['bank_name']}"

        attachment = {
            "ref": ref,
            "title": title,
            "file": filename,
            "signed_file": "",
            "status": "unsigned",
            "date_prepared": TODAY_SHORT,
            "date_sent": "",
            "date_signed": "",
        }

        results.append({
            "ref": ref,
            "folder": config["folder"],
            "filename": filename,
            "filepath": filepath,
            "wp": config["wp"],
            "attachment": attachment,
        })

        print(f"  [OK] {ref} -> {config['folder']}/{filename}")

    return results


def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_templates.py all")
        print("       python generate_templates.py T1 T4 T8")
        sys.exit(1)

    args = [a.upper() for a in sys.argv[1:] if not a.startswith("--")]

    # Determine templates to generate
    if "ALL" in args:
        templates = sorted(TEMPLATE_MAP.keys(), key=lambda x: int(x[1:]))
    else:
        templates = [a for a in args if a in TEMPLATE_MAP]
        if not templates:
            print(f"ERROR: No valid templates specified. Available: {', '.join(sorted(TEMPLATE_MAP.keys(), key=lambda x: int(x[1:])))}")
            sys.exit(1)

    print(f"\n{'='*60}")
    print(f"  AUDIT TEMPLATE GENERATOR")
    print(f"  Templates: {', '.join(templates)}")
    print(f"  Date: {TODAY}")
    print(f"{'='*60}\n")

    # Load master data
    vd = load_master_data()
    print(f"Loaded master_data.json ({len(vd)} variables)\n")

    all_results = []
    for ref in templates:
        print(f"Generating {ref} — {TEMPLATE_MAP[ref]['title']}...")
        results = generate_template(ref, vd)
        all_results.extend(results)

    # Output summary
    print(f"\n{'='*60}")
    print(f"  GENERATION COMPLETE: {len(all_results)} file(s)")
    print(f"{'='*60}\n")

    # Output ATTACHMENT JSON for Claude to process
    print("=== ATTACHMENT_COMMENTS ===")
    for r in all_results:
        if r["wp"]:
            print(json.dumps({
                "wp": r["wp"],
                "folder": r["folder"],
                "attachment_comment": f'<!-- ATTACHMENT:{json.dumps(r["attachment"])} -->',
            }))
    print("=== END_ATTACHMENT_COMMENTS ===")

    # Output summary JSON
    print("\n=== GENERATED_FILES ===")
    print(json.dumps([{
        "ref": r["ref"],
        "folder": r["folder"],
        "filename": r["filename"],
        "wp": r["wp"],
    } for r in all_results], indent=2))
    print("=== END_GENERATED_FILES ===")


if __name__ == "__main__":
    main()
