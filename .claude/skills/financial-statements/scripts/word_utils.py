"""
AUDIT WORD DOCUMENT UTILITIES
==============================
Python utilities for creating Word-based audit documents including:
- Draft Financial Statements
- Audit Reports
- Management Letters
- Representation Letters
Designed for Malaysian statutory audit engagements.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

# ============================================================================
# STYLE CONFIGURATIONS
# ============================================================================

def set_document_defaults(doc):
    """Set default document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Set default paragraph spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.15


def add_heading_style(doc, name, font_size, bold=True, space_before=12, space_after=6):
    """Add custom heading style."""
    styles = doc.styles
    try:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    except:
        style = styles[name]

    font = style.font
    font.name = 'Arial'
    font.size = Pt(font_size)
    font.bold = bold

    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

    return style


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def format_currency(value, show_brackets_negative=True):
    """Format number as currency."""
    if value is None:
        return "-"
    if isinstance(value, str):
        return value
    if value < 0 and show_brackets_negative:
        return f"({abs(value):,.0f})"
    return f"{value:,.0f}"


# ============================================================================
# FINANCIAL STATEMENTS TEMPLATES
# ============================================================================

def create_financial_statements(
    filepath: str,
    company_name: str,
    company_no: str,
    year_end: str,
    prior_year_end: str = None,
    reporting_framework: str = "MPERS",
    data: dict = None
):
    """
    Create complete draft financial statements package.

    Args:
        filepath: Output file path
        company_name: Legal name of company
        company_no: Company registration number
        year_end: Current financial year end
        prior_year_end: Prior year end (for comparatives)
        reporting_framework: MPERS or MFRS
        data: Dictionary containing financial data
    """
    doc = Document()
    set_document_defaults(doc)

    # -------------------------------------------------------------------------
    # COVER PAGE
    # -------------------------------------------------------------------------
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name.upper())
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"(Company No: {company_no})")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(Incorporated in Malaysia)")
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FINANCIAL STATEMENTS")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"FOR THE FINANCIAL YEAR ENDED {year_end.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # CONTENTS PAGE
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"(Company No: {company_no})")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONTENTS")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Contents table
    contents = [
        ("Directors' Report", "1 - 3"),
        ("Statement by Directors", "4"),
        ("Statutory Declaration", "5"),
        ("Independent Auditors' Report", "6 - 8"),
        ("Statement of Financial Position", "9"),
        ("Statement of Comprehensive Income", "10"),
        ("Statement of Changes in Equity", "11"),
        ("Statement of Cash Flows", "12"),
        ("Notes to the Financial Statements", "13 - XX"),
    ]

    table = doc.add_table(rows=len(contents), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, (item, page) in enumerate(contents):
        row = table.rows[idx]
        row.cells[0].text = item
        row.cells[1].text = page
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3)

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # DIRECTORS' REPORT
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"(Company No: {company_no})")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DIRECTORS' REPORT")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    doc.add_paragraph(
        f"The directors hereby submit their report together with the audited financial statements "
        f"of the Company for the financial year ended {year_end}."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("PRINCIPAL ACTIVITIES")
    run.bold = True

    doc.add_paragraph(
        "The principal activities of the Company are [DESCRIBE PRINCIPAL ACTIVITIES]. "
        "There have been no significant changes in the nature of these activities during the financial year."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("RESULTS")
    run.bold = True

    # Results table
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = ""
    table.rows[0].cells[1].text = "RM"
    table.rows[1].cells[0].text = "Profit/(Loss) for the financial year"
    table.rows[1].cells[1].text = "[AMOUNT]"

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("DIVIDENDS")
    run.bold = True

    doc.add_paragraph(
        "No dividend has been paid or declared by the Company since the end of the previous financial year. "
        "The directors do not recommend any dividend payment in respect of the current financial year."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("RESERVES AND PROVISIONS")
    run.bold = True

    doc.add_paragraph(
        "There were no material transfers to or from reserves or provisions during the financial year."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("DIRECTORS")
    run.bold = True

    doc.add_paragraph("The directors who served during the financial year until the date of this report are:")
    doc.add_paragraph("[DIRECTOR NAME 1]", style='List Bullet')
    doc.add_paragraph("[DIRECTOR NAME 2]", style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("DIRECTORS' INTERESTS")
    run.bold = True

    doc.add_paragraph(
        "According to the Register of Directors' Shareholdings required to be kept by the Company under "
        "Section 59 of the Companies Act 2016, the interests of directors in office at the end of the "
        "financial year in shares of the Company are as follows:"
    )

    doc.add_paragraph("[INSERT DIRECTORS' INTERESTS TABLE OR STATE NIL]")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("DIRECTORS' BENEFITS")
    run.bold = True

    doc.add_paragraph(
        "Since the end of the previous financial year, no director has received or become entitled to receive "
        "any benefit (other than benefits included in the aggregate amount of emoluments received or due and "
        "receivable by the directors as disclosed in Note XX to the financial statements) by reason of a contract "
        "made by the Company or a related corporation with the director or with a firm of which the director is "
        "a member, or with a company in which the director has a substantial financial interest."
    )

    doc.add_paragraph(
        "Neither during nor at the end of the financial year was the Company a party to any arrangement "
        "whose object was to enable the directors to acquire benefits through the acquisition of shares in "
        "or debentures of the Company or any other body corporate."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("INDEMNITY AND INSURANCE FOR DIRECTORS, OFFICERS AND AUDITORS")
    run.bold = True

    doc.add_paragraph(
        "The Company has not indemnified or agreed to indemnify any director, officer or auditor of the Company. "
        "During the financial year, there were no premiums paid for any insurance to indemnify directors, "
        "officers or auditors of the Company."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("AUDITORS")
    run.bold = True

    doc.add_paragraph(
        "The auditors, [AUDITOR FIRM NAME], have expressed their willingness to continue in office."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("AUDITORS' REMUNERATION")
    run.bold = True

    doc.add_paragraph(
        "The details of the auditors' remuneration are disclosed in Note XX to the financial statements."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(
        f"Signed on behalf of the Board of Directors in accordance with a resolution of the directors dated "
        f"[DATE]."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_____________________________")
    doc.add_paragraph("[DIRECTOR NAME]")
    doc.add_paragraph("Director")

    doc.add_paragraph()
    doc.add_paragraph("_____________________________")
    doc.add_paragraph("[DIRECTOR NAME]")
    doc.add_paragraph("Director")

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # STATEMENT OF FINANCIAL POSITION
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"(Company No: {company_no})")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATEMENT OF FINANCIAL POSITION")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"AS AT {year_end.upper()}")
    run.bold = True

    doc.add_paragraph()

    # Create SOFP table
    sofp_items = [
        ("", "Note", f"{year_end[:4]}", f"{int(year_end[:4])-1}" if prior_year_end else ""),
        ("", "", "RM", "RM"),
        ("ASSETS", "", "", ""),
        ("Non-Current Assets", "", "", ""),
        ("Property, plant and equipment", "5", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "", ""),
        ("Current Assets", "", "", ""),
        ("Trade receivables", "6", "[AMOUNT]", "[AMOUNT]"),
        ("Other receivables, deposits and prepayments", "7", "[AMOUNT]", "[AMOUNT]"),
        ("Cash and bank balances", "8", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "________", "________"),
        ("", "", "[TOTAL]", "[TOTAL]"),
        ("", "", "", ""),
        ("TOTAL ASSETS", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "========", "========"),
        ("", "", "", ""),
        ("EQUITY AND LIABILITIES", "", "", ""),
        ("Equity", "", "", ""),
        ("Share capital", "9", "[AMOUNT]", "[AMOUNT]"),
        ("Retained earnings/(Accumulated losses)", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "________", "________"),
        ("Total Equity", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "", ""),
        ("Non-Current Liabilities", "", "", ""),
        ("Amount due to holding company", "10", "[AMOUNT]", "[AMOUNT]"),
        ("Amount due to related companies", "11", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "________", "________"),
        ("", "", "[TOTAL]", "[TOTAL]"),
        ("", "", "", ""),
        ("Current Liabilities", "", "", ""),
        ("Trade payables", "12", "[AMOUNT]", "[AMOUNT]"),
        ("Other payables and accruals", "13", "[AMOUNT]", "[AMOUNT]"),
        ("Amount due to directors", "14", "[AMOUNT]", "[AMOUNT]"),
        ("Tax payable", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "________", "________"),
        ("", "", "[TOTAL]", "[TOTAL]"),
        ("", "", "", ""),
        ("Total Liabilities", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "________", "________"),
        ("", "", "", ""),
        ("TOTAL EQUITY AND LIABILITIES", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "========", "========"),
    ]

    table = doc.add_table(rows=len(sofp_items), cols=4)

    for row_idx, row_data in enumerate(sofp_items):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(cell_data)
            if col_idx >= 2:  # Amount columns
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if row_idx == 0:  # Header row
                cell.paragraphs[0].runs[0].bold = True if cell.text else False

    # Set column widths
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(0.5)
    table.columns[2].width = Inches(1.2)
    table.columns[3].width = Inches(1.2)

    doc.add_paragraph()
    doc.add_paragraph("The accompanying notes form an integral part of the financial statements.")

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # STATEMENT OF COMPREHENSIVE INCOME
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"(Company No: {company_no})")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATEMENT OF COMPREHENSIVE INCOME")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"FOR THE FINANCIAL YEAR ENDED {year_end.upper()}")
    run.bold = True

    doc.add_paragraph()

    soci_items = [
        ("", "Note", f"{year_end[:4]}", f"{int(year_end[:4])-1}" if prior_year_end else ""),
        ("", "", "RM", "RM"),
        ("Revenue", "15", "[AMOUNT]", "[AMOUNT]"),
        ("Cost of sales", "16", "([AMOUNT])", "([AMOUNT])"),
        ("", "", "________", "________"),
        ("Gross profit", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "", ""),
        ("Other income", "17", "[AMOUNT]", "[AMOUNT]"),
        ("Administrative expenses", "18", "([AMOUNT])", "([AMOUNT])"),
        ("Selling and distribution expenses", "", "([AMOUNT])", "([AMOUNT])"),
        ("", "", "________", "________"),
        ("Profit/(Loss) from operations", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "", ""),
        ("Finance costs", "19", "([AMOUNT])", "([AMOUNT])"),
        ("", "", "________", "________"),
        ("Profit/(Loss) before tax", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "", ""),
        ("Tax expense", "20", "([AMOUNT])", "([AMOUNT])"),
        ("", "", "________", "________"),
        ("Profit/(Loss) for the year", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "========", "========"),
        ("", "", "", ""),
        ("Other comprehensive income", "", "-", "-"),
        ("", "", "________", "________"),
        ("Total comprehensive income/(loss) for the year", "", "[AMOUNT]", "[AMOUNT]"),
        ("", "", "========", "========"),
    ]

    table = doc.add_table(rows=len(soci_items), cols=4)

    for row_idx, row_data in enumerate(soci_items):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(cell_data)
            if col_idx >= 2:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True if cell.text else False

    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(0.5)
    table.columns[2].width = Inches(1.2)
    table.columns[3].width = Inches(1.2)

    doc.add_paragraph()
    doc.add_paragraph("The accompanying notes form an integral part of the financial statements.")

    doc.add_page_break()

    # -------------------------------------------------------------------------
    # NOTES TO FINANCIAL STATEMENTS - TEMPLATE
    # -------------------------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{company_name.upper()}")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"(Company No: {company_no})")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOTES TO THE FINANCIAL STATEMENTS")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"FOR THE FINANCIAL YEAR ENDED {year_end.upper()}")
    run.bold = True

    doc.add_paragraph()

    # Note 1
    p = doc.add_paragraph()
    run = p.add_run("1. CORPORATE INFORMATION")
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        f"The Company is a private limited company, incorporated and domiciled in Malaysia. "
        f"The registered office of the Company is located at [REGISTERED ADDRESS]."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"The principal place of business of the Company is located at [BUSINESS ADDRESS]."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"The principal activities of the Company are [PRINCIPAL ACTIVITIES]. "
        f"There have been no significant changes in the nature of these activities during the financial year."
    )

    doc.add_paragraph()

    # Note 2
    p = doc.add_paragraph()
    run = p.add_run("2. BASIS OF PREPARATION")
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        f"The financial statements of the Company have been prepared in accordance with "
        f"{reporting_framework} and the requirements of the Companies Act 2016 in Malaysia."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "The financial statements have been prepared on the historical cost basis unless otherwise "
        "indicated in the summary of significant accounting policies."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "The financial statements are presented in Ringgit Malaysia (RM), which is the Company's "
        "functional currency."
    )

    doc.add_paragraph()

    # Note 3
    p = doc.add_paragraph()
    run = p.add_run("3. SIGNIFICANT ACCOUNTING POLICIES")
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("[INSERT ACCOUNTING POLICIES - PPE, FINANCIAL INSTRUMENTS, REVENUE, ETC.]")

    # Save document
    doc.save(filepath)
    print(f"Financial statements created: {filepath}")
    return filepath


def create_management_letter(
    filepath: str,
    company_name: str,
    year_end: str,
    findings: list,
    auditor_firm: str,
    addressee: str = "Board of Directors"
):
    """
    Create management letter with audit findings.

    Args:
        findings: List of dicts with keys: title, observation, risk, recommendation, management_response
    """
    doc = Document()
    set_document_defaults(doc)

    # Letterhead placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[{auditor_firm.upper()}]")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("[FIRM ADDRESS]")
    doc.add_paragraph()

    # Date
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph()

    # Addressee
    doc.add_paragraph(f"The {addressee}")
    doc.add_paragraph(company_name)
    doc.add_paragraph("[COMPANY ADDRESS]")
    doc.add_paragraph()

    doc.add_paragraph(f"Dear Sirs,")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(f"MANAGEMENT LETTER FOR THE FINANCIAL YEAR ENDED {year_end.upper()}")
    run.bold = True
    run.underline = True

    doc.add_paragraph()
    doc.add_paragraph(
        f"We have completed our audit of the financial statements of {company_name} for the "
        f"financial year ended {year_end}. During the course of our audit, we noted certain matters "
        f"relating to internal controls and other operational matters which we would like to bring to "
        f"your attention."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "This letter does not affect our report dated [DATE] on the financial statements of the Company."
    )

    doc.add_paragraph()

    # Findings
    for idx, finding in enumerate(findings, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {finding.get('title', 'Finding')}")
        run.bold = True

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Observation:")
        run.bold = True
        doc.add_paragraph(finding.get('observation', '[Description of observation]'))

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Risk/Implication:")
        run.bold = True
        doc.add_paragraph(finding.get('risk', '[Description of risk]'))

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Recommendation:")
        run.bold = True
        doc.add_paragraph(finding.get('recommendation', '[Recommendation]'))

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Management Response:")
        run.bold = True
        doc.add_paragraph(finding.get('management_response', '[To be completed by management]'))

        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph()

    # Closing
    doc.add_paragraph(
        "We would like to take this opportunity to thank the management and staff for their "
        "co-operation and assistance rendered during the course of our audit."
    )

    doc.add_paragraph()
    doc.add_paragraph("Please do not hesitate to contact us if you require any clarification on the above matters.")
    doc.add_paragraph()
    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph(f"_____________________________")
    doc.add_paragraph(f"{auditor_firm.upper()}")
    doc.add_paragraph("AF: [XXXX]")

    doc.save(filepath)
    print(f"Management letter created: {filepath}")
    return filepath


def create_representation_letter(
    filepath: str,
    company_name: str,
    company_no: str,
    year_end: str,
    auditor_firm: str,
    directors: list
):
    """
    Create management representation letter template.
    """
    doc = Document()
    set_document_defaults(doc)

    # Company letterhead placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[{company_name.upper()} LETTERHEAD]")
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_paragraph(f"Date: [DATE]")
    doc.add_paragraph()

    doc.add_paragraph(auditor_firm)
    doc.add_paragraph("[AUDITOR ADDRESS]")
    doc.add_paragraph()

    doc.add_paragraph("Dear Sirs,")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(f"MANAGEMENT REPRESENTATION LETTER")
    run.bold = True
    run.underline = True

    p = doc.add_paragraph()
    run = p.add_run(f"AUDIT OF FINANCIAL STATEMENTS FOR THE FINANCIAL YEAR ENDED {year_end.upper()}")
    run.bold = True
    run.underline = True

    doc.add_paragraph()

    doc.add_paragraph(
        f"This representation letter is provided in connection with your audit of the financial statements "
        f"of {company_name} (Company No: {company_no}) for the financial year ended {year_end} for the "
        f"purpose of expressing an opinion as to whether the financial statements give a true and fair view "
        f"in accordance with Malaysian Private Entities Reporting Standard and the requirements of the "
        f"Companies Act 2016 in Malaysia."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "We confirm that, to the best of our knowledge and belief, having made such inquiries as we "
        "considered necessary for the purpose of appropriately informing ourselves:"
    )

    doc.add_paragraph()

    # Representations
    representations = [
        "We have fulfilled our responsibilities, as set out in the terms of the audit engagement, for the preparation of the financial statements in accordance with Malaysian Private Entities Reporting Standard and the requirements of the Companies Act 2016 in Malaysia.",
        "Significant assumptions used by us in making accounting estimates, including those measured at fair value, are reasonable.",
        "Related party relationships and transactions have been appropriately accounted for and disclosed in accordance with the requirements of the applicable financial reporting framework.",
        "All events subsequent to the date of the financial statements and for which the applicable financial reporting framework requires adjustment or disclosure have been adjusted or disclosed.",
        "The effects of uncorrected misstatements are immaterial, both individually and in the aggregate, to the financial statements as a whole.",
        "We have provided you with all relevant information and access as agreed in the terms of the audit engagement.",
        "All transactions have been recorded in the accounting records and are reflected in the financial statements.",
        "We have disclosed to you the results of our assessment of the risk that the financial statements may be materially misstated as a result of fraud.",
        "We have disclosed to you all information in relation to fraud or suspected fraud that we are aware of and that affects the Company.",
        "We have disclosed to you all information in relation to allegations of fraud, or suspected fraud, affecting the Company's financial statements communicated by employees, former employees, analysts, regulators or others.",
        "We have disclosed to you all known instances of non-compliance or suspected non-compliance with laws and regulations whose effects should be considered when preparing financial statements.",
        "We have disclosed to you the identity of the Company's related parties and all the related party relationships and transactions of which we are aware.",
        "We have no plans or intentions that may materially affect the carrying values or classification of assets and liabilities reflected in the financial statements.",
        "The Company has satisfactory title to all assets, and there are no liens or encumbrances on such assets, except as disclosed in the notes to the financial statements.",
        "We have recorded or disclosed, as appropriate, all liabilities, both actual and contingent.",
    ]

    for idx, rep in enumerate(representations, 1):
        doc.add_paragraph(f"{idx}. {rep}")

    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_paragraph("Yours faithfully,")
    doc.add_paragraph()

    for director in directors:
        doc.add_paragraph()
        doc.add_paragraph("_____________________________")
        doc.add_paragraph(f"{director}")
        doc.add_paragraph("Director")
        doc.add_paragraph()

    doc.save(filepath)
    print(f"Representation letter created: {filepath}")
    return filepath


if __name__ == "__main__":
    print("Audit Word Document Utilities loaded successfully.")
    print("Available functions:")
    print("  - create_financial_statements()")
    print("  - create_management_letter()")
    print("  - create_representation_letter()")
