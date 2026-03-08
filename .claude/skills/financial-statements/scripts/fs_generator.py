"""
MPERS FINANCIAL STATEMENTS GENERATOR
=====================================
Generate complete MPERS-compliant financial statements in Word format.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


class MPERSFinancialStatements:
    """
    Generate complete MPERS-compliant financial statements.
    """

    def __init__(
        self,
        company_name: str,
        company_no: str,
        year_end: str,
        year_end_date: str = None,
        prior_year_end: str = None,
        registered_address: str = None,
        business_address: str = None,
        principal_activities: str = None,
        directors: list = None,
        auditor_firm: str = None,
        auditor_af_no: str = None
    ):
        self.company_name = company_name
        self.company_no = company_no
        self.year_end = year_end
        self.year_end_date = year_end_date or year_end
        self.prior_year_end = prior_year_end
        self.registered_address = registered_address or "[REGISTERED ADDRESS]"
        self.business_address = business_address or "[BUSINESS ADDRESS]"
        self.principal_activities = principal_activities or "[DESCRIBE PRINCIPAL ACTIVITIES]"
        self.directors = directors or ["[DIRECTOR 1]", "[DIRECTOR 2]"]
        self.auditor_firm = auditor_firm or "[AUDITOR FIRM NAME]"
        self.auditor_af_no = auditor_af_no or "[AF XXXX]"

        # Financial data
        self.sofp_data = {}
        self.soci_data = {}
        self.notes_data = {}

        # Initialize document
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup document styles."""
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Add custom styles
        styles = self.doc.styles

        # Title style
        try:
            title_style = styles.add_style('FSTitle', WD_STYLE_TYPE.PARAGRAPH)
        except:
            title_style = styles['FSTitle']
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(14)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(6)

        # Subtitle style
        try:
            subtitle_style = styles.add_style('FSSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        except:
            subtitle_style = styles['FSSubtitle']
        subtitle_style.font.name = 'Arial'
        subtitle_style.font.size = Pt(12)
        subtitle_style.font.bold = True
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(12)

        # Note heading style
        try:
            note_style = styles.add_style('NoteHeading', WD_STYLE_TYPE.PARAGRAPH)
        except:
            note_style = styles['NoteHeading']
        note_style.font.name = 'Arial'
        note_style.font.size = Pt(11)
        note_style.font.bold = True
        note_style.paragraph_format.space_before = Pt(12)
        note_style.paragraph_format.space_after = Pt(6)

    def _add_company_header(self):
        """Add standard company header."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.company_name.upper())
        run.bold = True
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"(Company No: {self.company_no})")
        run.font.size = Pt(11)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("(Incorporated in Malaysia)")
        run.font.size = Pt(11)

        self.doc.add_paragraph()

    def set_sofp_data(self, data: dict):
        """Set Statement of Financial Position data."""
        self.sofp_data = data

    def set_soci_data(self, data: dict):
        """Set Statement of Comprehensive Income data."""
        self.soci_data = data

    def set_notes_data(self, data: dict):
        """Set Notes data."""
        self.notes_data = data

    def generate_cover_page(self):
        """Generate cover page."""
        # Add spacing
        for _ in range(8):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.company_name.upper())
        run.bold = True
        run.font.size = Pt(20)

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"(Company No: {self.company_no})")
        run.font.size = Pt(12)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("(Incorporated in Malaysia)")
        run.font.size = Pt(12)

        for _ in range(4):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("FINANCIAL STATEMENTS")
        run.bold = True
        run.font.size = Pt(18)

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"FOR THE FINANCIAL YEAR ENDED")
        run.bold = True
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self.year_end.upper()}")
        run.bold = True
        run.font.size = Pt(14)

        self.doc.add_page_break()

    def generate_contents(self):
        """Generate contents page."""
        self._add_company_header()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONTENTS")
        run.bold = True
        run.font.size = Pt(14)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        contents = [
            ("Directors' Report", "1 - 4"),
            ("Statement by Directors", "5"),
            ("Statutory Declaration", "6"),
            ("Independent Auditors' Report", "7 - 9"),
            ("Statement of Financial Position", "10"),
            ("Statement of Comprehensive Income", "11"),
            ("Statement of Changes in Equity", "12"),
            ("Statement of Cash Flows", "13"),
            ("Notes to the Financial Statements", "14 - XX"),
        ]

        table = self.doc.add_table(rows=len(contents), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx, (item, page) in enumerate(contents):
            row = table.rows[idx]
            cell1 = row.cells[0]
            cell2 = row.cells[1]

            cell1.text = item
            cell2.text = page
            cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Set widths
            cell1.width = Inches(4.5)
            cell2.width = Inches(1)

        self.doc.add_page_break()

    def generate_directors_report(self):
        """Generate Directors' Report."""
        self._add_company_header()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DIRECTORS' REPORT")
        run.bold = True
        run.font.size = Pt(14)

        self.doc.add_paragraph()

        self.doc.add_paragraph(
            f"The directors hereby submit their report together with the audited financial "
            f"statements of the Company for the financial year ended {self.year_end}."
        )

        self.doc.add_paragraph()

        # Principal Activities
        p = self.doc.add_paragraph()
        run = p.add_run("PRINCIPAL ACTIVITIES")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            f"The principal activities of the Company are {self.principal_activities}. "
            f"There have been no significant changes in the nature of these activities "
            f"during the financial year."
        )

        self.doc.add_paragraph()

        # Results
        p = self.doc.add_paragraph()
        run = p.add_run("RESULTS")
        run.bold = True

        self.doc.add_paragraph()

        table = self.doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = ""
        table.rows[0].cells[1].text = "RM"
        table.rows[1].cells[0].text = "Profit/(Loss) for the financial year"
        profit = self.soci_data.get('profit_for_year', '[AMOUNT]')
        table.rows[1].cells[1].text = f"{profit:,.0f}" if isinstance(profit, (int, float)) else str(profit)

        self.doc.add_paragraph()

        # Dividends
        p = self.doc.add_paragraph()
        run = p.add_run("DIVIDENDS")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "No dividend has been paid or declared by the Company since the end of the "
            "previous financial year. The directors do not recommend any dividend payment "
            "in respect of the current financial year."
        )

        self.doc.add_paragraph()

        # Reserves and Provisions
        p = self.doc.add_paragraph()
        run = p.add_run("RESERVES AND PROVISIONS")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "There were no material transfers to or from reserves or provisions during "
            "the financial year."
        )

        self.doc.add_paragraph()

        # Directors
        p = self.doc.add_paragraph()
        run = p.add_run("DIRECTORS")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The directors who served during the financial year until the date of this report are:"
        )

        for director in self.directors:
            self.doc.add_paragraph(director, style='List Bullet')

        self.doc.add_paragraph()

        # Directors' Interests
        p = self.doc.add_paragraph()
        run = p.add_run("DIRECTORS' INTERESTS")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "According to the Register of Directors' Shareholdings required to be kept by the "
            "Company under Section 59 of the Companies Act 2016, the interests of directors in "
            "office at the end of the financial year in shares of the Company are as follows:"
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph("[INSERT DIRECTORS' INTERESTS TABLE OR STATE NIL]")

        self.doc.add_paragraph()

        # Directors' Benefits
        p = self.doc.add_paragraph()
        run = p.add_run("DIRECTORS' BENEFITS")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Since the end of the previous financial year, no director has received or become "
            "entitled to receive any benefit (other than benefits included in the aggregate amount "
            "of emoluments received or due and receivable by the directors as disclosed in Note [X] "
            "to the financial statements) by reason of a contract made by the Company or a related "
            "corporation with the director or with a firm of which the director is a member, or with "
            "a company in which the director has a substantial financial interest."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Neither during nor at the end of the financial year was the Company a party to any "
            "arrangement whose object was to enable the directors to acquire benefits through the "
            "acquisition of shares in or debentures of the Company or any other body corporate."
        )

        self.doc.add_paragraph()

        # Indemnity and Insurance
        p = self.doc.add_paragraph()
        run = p.add_run("INDEMNITY AND INSURANCE FOR DIRECTORS, OFFICERS AND AUDITORS")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The Company has not indemnified or agreed to indemnify any director, officer or "
            "auditor of the Company. During the financial year, there were no premiums paid for "
            "any insurance to indemnify directors, officers or auditors of the Company."
        )

        self.doc.add_paragraph()

        # Other Statutory Information
        p = self.doc.add_paragraph()
        run = p.add_run("OTHER STATUTORY INFORMATION")
        run.bold = True

        self.doc.add_paragraph()

        self.doc.add_paragraph(
            "(a) Before the financial statements of the Company were prepared, the directors took "
            "reasonable steps:"
        )

        self.doc.add_paragraph(
            "    (i) to ascertain that proper action had been taken in relation to the writing off "
            "of bad debts and the making of provision for doubtful debts and satisfied themselves "
            "that all known bad debts had been written off and that adequate provision had been "
            "made for doubtful debts; and"
        )

        self.doc.add_paragraph(
            "    (ii) to ensure that any current assets which were unlikely to be realised their "
            "values as shown in the accounting records in the ordinary course of business had been "
            "written down to an amount which they might be expected so to realise."
        )

        self.doc.add_paragraph()

        self.doc.add_paragraph(
            "(b) At the date of this report, the directors are not aware of any circumstances:"
        )

        self.doc.add_paragraph(
            "    (i) which would render the amounts written off for bad debts or the amount of the "
            "provision for doubtful debts in the financial statements of the Company inadequate to "
            "any substantial extent; or"
        )

        self.doc.add_paragraph(
            "    (ii) which would render the values attributed to current assets in the financial "
            "statements of the Company misleading; or"
        )

        self.doc.add_paragraph(
            "    (iii) which have arisen which would render adherence to the existing method of "
            "valuation of assets or liabilities of the Company misleading or inappropriate."
        )

        self.doc.add_paragraph()

        self.doc.add_paragraph("(c) At the date of this report:")

        self.doc.add_paragraph(
            "    (i) there are no charges on the assets of the Company which have arisen since the "
            "end of the financial year to secure the liabilities of any other person; and"
        )

        self.doc.add_paragraph(
            "    (ii) there are no contingent liabilities which have arisen since the end of the "
            "financial year."
        )

        self.doc.add_paragraph()

        self.doc.add_paragraph("(d) In the opinion of the directors:")

        self.doc.add_paragraph(
            "    (i) no contingent liability or other liability has become enforceable or is likely "
            "to become enforceable within the period of twelve months after the end of the financial "
            "year which will or may affect the ability of the Company to meet its obligations as and "
            "when they fall due;"
        )

        self.doc.add_paragraph(
            "    (ii) the results of the operations of the Company during the financial year were "
            "not substantially affected by any item, transaction or event of a material and unusual "
            "nature; and"
        )

        self.doc.add_paragraph(
            "    (iii) there has not arisen in the interval between the end of the financial year "
            "and the date of this report any item, transaction or event of a material and unusual "
            "nature likely to affect substantially the results of the operations of the Company "
            "for the financial year in which this report is made."
        )

        self.doc.add_paragraph()

        # Auditors
        p = self.doc.add_paragraph()
        run = p.add_run("AUDITORS")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            f"The auditors, {self.auditor_firm}, have expressed their willingness to continue in office."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The details of the auditors' remuneration are disclosed in Note [X] to the financial statements."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        self.doc.add_paragraph(
            f"Signed on behalf of the Board of Directors in accordance with a resolution of the "
            f"directors dated [DATE]."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Signature lines
        table = self.doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = "_____________________________"
        table.rows[0].cells[1].text = "_____________________________"
        table.rows[1].cells[0].text = self.directors[0] if len(self.directors) > 0 else "[DIRECTOR NAME]"
        table.rows[1].cells[1].text = self.directors[1] if len(self.directors) > 1 else "[DIRECTOR NAME]"
        table.rows[2].cells[0].text = "Director"
        table.rows[2].cells[1].text = "Director"

        self.doc.add_paragraph()
        self.doc.add_paragraph("[PLACE]")

        self.doc.add_page_break()

    def generate_statement_by_directors(self):
        """Generate Statement by Directors."""
        self._add_company_header()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("STATEMENT BY DIRECTORS")
        run.bold = True
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Pursuant to Section 251(2) of the Companies Act 2016")
        run.font.size = Pt(11)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        dir1 = self.directors[0] if len(self.directors) > 0 else "[DIRECTOR NAME 1]"
        dir2 = self.directors[1] if len(self.directors) > 1 else "[DIRECTOR NAME 2]"

        self.doc.add_paragraph(
            f"We, {dir1.upper()} and {dir2.upper()}, being two of the directors of "
            f"{self.company_name.upper()}, do hereby state that, in the opinion of the directors, "
            f"the financial statements set out on pages [X] to [X] are drawn up in accordance with "
            f"Malaysian Private Entities Reporting Standard and the requirements of the Companies "
            f"Act 2016 in Malaysia so as to give a true and fair view of the financial position of "
            f"the Company as at {self.year_end} and of its financial performance and cash flows for "
            f"the financial year then ended."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        self.doc.add_paragraph(
            f"Signed on behalf of the Board of Directors in accordance with a resolution of the "
            f"directors dated [DATE]."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        # Signature lines
        table = self.doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = "_____________________________"
        table.rows[0].cells[1].text = "_____________________________"
        table.rows[1].cells[0].text = dir1
        table.rows[1].cells[1].text = dir2
        table.rows[2].cells[0].text = "Director"
        table.rows[2].cells[1].text = "Director"

        self.doc.add_paragraph()
        self.doc.add_paragraph("[PLACE]")

        self.doc.add_page_break()

    def generate_statutory_declaration(self):
        """Generate Statutory Declaration."""
        self._add_company_header()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("STATUTORY DECLARATION")
        run.bold = True
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Pursuant to Section 251(1)(b) of the Companies Act 2016")
        run.font.size = Pt(11)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        self.doc.add_paragraph(
            f"I, [NAME], being the [POSITION] primarily responsible for the financial management "
            f"of {self.company_name.upper()}, do solemnly and sincerely declare that to the best "
            f"of my knowledge and belief, the financial statements set out on pages [X] to [X] are "
            f"correct and I make this solemn declaration conscientiously believing the same to be "
            f"true and by virtue of the provisions of the Statutory Declarations Act 1960."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        self.doc.add_paragraph("Subscribed and solemnly declared by")
        self.doc.add_paragraph("the abovenamed at [PLACE] in the")
        self.doc.add_paragraph("state of [STATE] on [DATE]")

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("_____________________________")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("[NAME]")

        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        self.doc.add_paragraph("Before me,")
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph("_____________________________")
        self.doc.add_paragraph("Commissioner for Oaths")

        self.doc.add_page_break()

    def generate_sofp(self):
        """Generate Statement of Financial Position."""
        self._add_company_header()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("STATEMENT OF FINANCIAL POSITION")
        run.bold = True
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"AS AT {self.year_end.upper()}")
        run.bold = True

        self.doc.add_paragraph()

        # Get data or use placeholders
        d = self.sofp_data

        def fmt(val):
            if val is None or val == '':
                return '-'
            if isinstance(val, (int, float)):
                if val == 0:
                    return '-'
                return f"{val:,.0f}"
            return str(val)

        # Create SOFP content
        sofp_lines = [
            ("", "Note", "2025", "2024"),
            ("", "", "RM", "RM"),
            ("ASSETS", "", "", ""),
            ("", "", "", ""),
            ("Non-Current Assets", "", "", ""),
            ("Property, plant and equipment", "5", fmt(d.get('ppe', '[XXX]')), fmt(d.get('ppe_py', '-'))),
            ("", "", "", ""),
            ("Current Assets", "", "", ""),
            ("Trade receivables", "6", fmt(d.get('trade_rec', '[XXX]')), fmt(d.get('trade_rec_py', '-'))),
            ("Other receivables, deposits and prepayments", "7", fmt(d.get('other_rec', '[XXX]')), fmt(d.get('other_rec_py', '-'))),
            ("Cash and bank balances", "8", fmt(d.get('cash', '[XXX]')), fmt(d.get('cash_py', '-'))),
            ("", "", "________", "________"),
            ("", "", fmt(d.get('total_ca', '[XXX]')), fmt(d.get('total_ca_py', '-'))),
            ("", "", "", ""),
            ("TOTAL ASSETS", "", fmt(d.get('total_assets', '[XXX]')), fmt(d.get('total_assets_py', '-'))),
            ("", "", "========", "========"),
            ("", "", "", ""),
            ("EQUITY AND LIABILITIES", "", "", ""),
            ("", "", "", ""),
            ("Equity", "", "", ""),
            ("Share capital", "9", fmt(d.get('share_cap', '[XXX]')), fmt(d.get('share_cap_py', '-'))),
            ("Retained earnings/(Accumulated losses)", "", fmt(d.get('retained', '[XXX]')), fmt(d.get('retained_py', '-'))),
            ("", "", "________", "________"),
            ("Total Equity", "", fmt(d.get('total_equity', '[XXX]')), fmt(d.get('total_equity_py', '-'))),
            ("", "", "", ""),
            ("Non-Current Liabilities", "", "", ""),
            ("Amount due to holding company", "10", fmt(d.get('due_holding_ncl', '[XXX]')), fmt(d.get('due_holding_ncl_py', '-'))),
            ("Amount due to related companies", "11", fmt(d.get('due_related_ncl', '[XXX]')), fmt(d.get('due_related_ncl_py', '-'))),
            ("", "", "________", "________"),
            ("", "", fmt(d.get('total_ncl', '[XXX]')), fmt(d.get('total_ncl_py', '-'))),
            ("", "", "", ""),
            ("Current Liabilities", "", "", ""),
            ("Trade payables", "12", fmt(d.get('trade_pay', '[XXX]')), fmt(d.get('trade_pay_py', '-'))),
            ("Other payables and accruals", "13", fmt(d.get('other_pay', '[XXX]')), fmt(d.get('other_pay_py', '-'))),
            ("Amount due to directors", "14", fmt(d.get('due_directors', '[XXX]')), fmt(d.get('due_directors_py', '-'))),
            ("Tax payable", "", fmt(d.get('tax_pay', '[XXX]')), fmt(d.get('tax_pay_py', '-'))),
            ("", "", "________", "________"),
            ("", "", fmt(d.get('total_cl', '[XXX]')), fmt(d.get('total_cl_py', '-'))),
            ("", "", "", ""),
            ("Total Liabilities", "", fmt(d.get('total_liab', '[XXX]')), fmt(d.get('total_liab_py', '-'))),
            ("", "", "________", "________"),
            ("", "", "", ""),
            ("TOTAL EQUITY AND LIABILITIES", "", fmt(d.get('total_eq_liab', '[XXX]')), fmt(d.get('total_eq_liab_py', '-'))),
            ("", "", "========", "========"),
        ]

        table = self.doc.add_table(rows=len(sofp_lines), cols=4)

        for row_idx, row_data in enumerate(sofp_lines):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_data)
                # Right align amounts
                if col_idx >= 2:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(3.5)
            row.cells[1].width = Inches(0.5)
            row.cells[2].width = Inches(1.2)
            row.cells[3].width = Inches(1.2)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("The accompanying notes form an integral part of the financial statements.")
        run.italic = True

        self.doc.add_page_break()

    def generate_soci(self):
        """Generate Statement of Comprehensive Income."""
        self._add_company_header()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("STATEMENT OF COMPREHENSIVE INCOME")
        run.bold = True
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"FOR THE FINANCIAL YEAR ENDED {self.year_end.upper()}")
        run.bold = True

        self.doc.add_paragraph()

        d = self.soci_data

        def fmt(val):
            if val is None or val == '':
                return '-'
            if isinstance(val, (int, float)):
                if val == 0:
                    return '-'
                if val < 0:
                    return f"({abs(val):,.0f})"
                return f"{val:,.0f}"
            return str(val)

        soci_lines = [
            ("", "Note", "2025", "2024"),
            ("", "", "RM", "RM"),
            ("", "", "", ""),
            ("Revenue", "15", fmt(d.get('revenue', '[XXX]')), fmt(d.get('revenue_py', '-'))),
            ("", "", "", ""),
            ("Cost of sales", "16", fmt(d.get('cos', '[XXX]')), fmt(d.get('cos_py', '-'))),
            ("", "", "________", "________"),
            ("Gross profit/(loss)", "", fmt(d.get('gross_profit', '[XXX]')), fmt(d.get('gross_profit_py', '-'))),
            ("", "", "", ""),
            ("Other income", "17", fmt(d.get('other_income', '[XXX]')), fmt(d.get('other_income_py', '-'))),
            ("", "", "", ""),
            ("Administrative expenses", "18", fmt(d.get('admin_exp', '[XXX]')), fmt(d.get('admin_exp_py', '-'))),
            ("", "", "", ""),
            ("Selling and distribution expenses", "", fmt(d.get('selling_exp', '[XXX]')), fmt(d.get('selling_exp_py', '-'))),
            ("", "", "________", "________"),
            ("Profit/(Loss) from operations", "", fmt(d.get('op_profit', '[XXX]')), fmt(d.get('op_profit_py', '-'))),
            ("", "", "", ""),
            ("Finance costs", "19", fmt(d.get('finance_cost', '[XXX]')), fmt(d.get('finance_cost_py', '-'))),
            ("", "", "________", "________"),
            ("Profit/(Loss) before tax", "20", fmt(d.get('pbt', '[XXX]')), fmt(d.get('pbt_py', '-'))),
            ("", "", "", ""),
            ("Tax expense", "21", fmt(d.get('tax_exp', '[XXX]')), fmt(d.get('tax_exp_py', '-'))),
            ("", "", "________", "________"),
            ("Profit/(Loss) for the year", "", fmt(d.get('profit_for_year', '[XXX]')), fmt(d.get('profit_for_year_py', '-'))),
            ("", "", "========", "========"),
            ("", "", "", ""),
            ("Other comprehensive income", "", "-", "-"),
            ("", "", "________", "________"),
            ("Total comprehensive income/(loss) for the year", "", fmt(d.get('total_ci', '[XXX]')), fmt(d.get('total_ci_py', '-'))),
            ("", "", "========", "========"),
        ]

        table = self.doc.add_table(rows=len(soci_lines), cols=4)

        for row_idx, row_data in enumerate(soci_lines):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_data)
                if col_idx >= 2:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for row in table.rows:
            row.cells[0].width = Inches(3.5)
            row.cells[1].width = Inches(0.5)
            row.cells[2].width = Inches(1.2)
            row.cells[3].width = Inches(1.2)

        self.doc.add_paragraph()
        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("The accompanying notes form an integral part of the financial statements.")
        run.italic = True

        self.doc.add_page_break()

    def generate_notes_template(self):
        """Generate Notes to Financial Statements template."""
        self._add_company_header()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("NOTES TO THE FINANCIAL STATEMENTS")
        run.bold = True
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"FOR THE FINANCIAL YEAR ENDED {self.year_end.upper()}")
        run.bold = True

        self.doc.add_paragraph()

        # Note 1 - Corporate Information
        p = self.doc.add_paragraph()
        run = p.add_run("1.  CORPORATE INFORMATION")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            f"The Company is a private limited liability company, incorporated and domiciled in Malaysia."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph(f"The registered office of the Company is located at:")
        self.doc.add_paragraph(self.registered_address)

        self.doc.add_paragraph()
        self.doc.add_paragraph(f"The principal place of business of the Company is located at:")
        self.doc.add_paragraph(self.business_address)

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            f"The principal activities of the Company are {self.principal_activities}. "
            f"There have been no significant changes in the nature of these activities during the financial year."
        )

        self.doc.add_paragraph()

        # Note 2 - Basis of Preparation
        p = self.doc.add_paragraph()
        run = p.add_run("2.  BASIS OF PREPARATION")
        run.bold = True

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        run = p.add_run("(a) Statement of compliance")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The financial statements of the Company have been prepared in accordance with Malaysian "
            "Private Entities Reporting Standard (\"MPERS\") and the requirements of the Companies Act "
            "2016 in Malaysia."
        )

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        run = p.add_run("(b) Basis of measurement")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The financial statements have been prepared on the historical cost basis, unless otherwise "
            "indicated in the significant accounting policies below."
        )

        self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        run = p.add_run("(c) Functional and presentation currency")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The financial statements are presented in Ringgit Malaysia (\"RM\"), which is the Company's "
            "functional currency. All financial information presented in RM has been rounded to the nearest "
            "RM, unless otherwise stated."
        )

        self.doc.add_paragraph()

        # Note 3 - Significant Accounting Policies
        p = self.doc.add_paragraph()
        run = p.add_run("3.  SIGNIFICANT ACCOUNTING POLICIES")
        run.bold = True

        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "The accounting policies set out below have been applied consistently to the periods presented "
            "in these financial statements."
        )

        self.doc.add_paragraph()
        self.doc.add_paragraph("[INSERT ACCOUNTING POLICIES AS APPLICABLE]")

        self.doc.add_paragraph()

        # Placeholder for remaining notes
        notes_list = [
            "4.  PROPERTY, PLANT AND EQUIPMENT",
            "5.  TRADE RECEIVABLES",
            "6.  OTHER RECEIVABLES, DEPOSITS AND PREPAYMENTS",
            "7.  CASH AND BANK BALANCES",
            "8.  SHARE CAPITAL",
            "9.  AMOUNT DUE TO HOLDING COMPANY",
            "10. AMOUNT DUE TO RELATED COMPANIES",
            "11. TRADE PAYABLES",
            "12. OTHER PAYABLES AND ACCRUALS",
            "13. AMOUNT DUE TO DIRECTORS",
            "14. REVENUE",
            "15. COST OF SALES",
            "16. OTHER INCOME",
            "17. ADMINISTRATIVE EXPENSES",
            "18. FINANCE COSTS",
            "19. TAX EXPENSE",
            "20. RELATED PARTY TRANSACTIONS",
            "21. CONTINGENT LIABILITIES",
            "22. CAPITAL COMMITMENTS",
        ]

        for note in notes_list:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run(note)
            run.bold = True
            self.doc.add_paragraph()
            self.doc.add_paragraph("[INSERT NOTE DETAILS]")

    def generate_full_fs(self, filepath: str):
        """Generate complete financial statements package."""
        self.generate_cover_page()
        self.generate_contents()
        self.generate_directors_report()
        self.generate_statement_by_directors()
        self.generate_statutory_declaration()
        self.generate_sofp()
        self.generate_soci()
        self.generate_notes_template()

        self.doc.save(filepath)
        print(f"Financial statements generated: {filepath}")
        return filepath


def generate_mpers_fs(
    filepath: str,
    company_name: str,
    company_no: str,
    year_end: str,
    sofp_data: dict = None,
    soci_data: dict = None,
    directors: list = None,
    auditor_firm: str = None,
    principal_activities: str = None
):
    """
    Convenience function to generate MPERS financial statements.
    """
    fs = MPERSFinancialStatements(
        company_name=company_name,
        company_no=company_no,
        year_end=year_end,
        directors=directors,
        auditor_firm=auditor_firm,
        principal_activities=principal_activities
    )

    if sofp_data:
        fs.set_sofp_data(sofp_data)
    if soci_data:
        fs.set_soci_data(soci_data)

    return fs.generate_full_fs(filepath)


if __name__ == "__main__":
    print("MPERS Financial Statements Generator")
    print("=====================================")
    print()
    print("Usage:")
    print("  from fs_generator import generate_mpers_fs, MPERSFinancialStatements")
    print()
    print("  # Quick generation")
    print("  generate_mpers_fs(")
    print("      filepath='output.docx',")
    print("      company_name='ABC Sdn. Bhd.',")
    print("      company_no='123456-X',")
    print("      year_end='31 December 2024'")
    print("  )")
    print()
    print("  # With data")
    print("  fs = MPERSFinancialStatements(...)")
    print("  fs.set_sofp_data({...})")
    print("  fs.set_soci_data({...})")
    print("  fs.generate_full_fs('output.docx')")
